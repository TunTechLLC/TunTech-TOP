"""QA-4 Revision Agent — tests for the deterministic matcher and in-place
docx applier (the genuinely risky code). The Claude judgment call
(generate_revision_edits) is validated by the E004 smoke test, consistent with
how the QA-1/2/3 detection calls are validated.
"""
from docx import Document

from api.services.qa_revision import (
    build_index,
    _para_index_at,
    locate_anchor,
    replace_in_paragraph,
    insert_paragraph_after,
    iter_paragraphs,
    reconcile_unaddressed,
)


# ---------------------------------------------------------------------------
# Reconciliation — every accepted item must be applied or flagged
# ---------------------------------------------------------------------------

def test_reconcile_unaddressed_flags_items_with_no_edit():
    accepted = [
        {'id': 'QH001', 'source': 'coherence', 'summary': 'mislabel'},
        {'id': 'QE004', 'source': 'editorial', 'summary': 'define PMO'},
        {'id': 'QC003', 'source': 'coverage', 'summary': 'time-sensitive item'},
    ]
    edits = [
        {'source_item_id': 'QH001'},
        {'source_item_id': 'QE004'},
        # QC003 has no edit — must be reported unaddressed
    ]
    missing = reconcile_unaddressed(accepted, edits)
    assert [m['id'] for m in missing] == ['QC003']


def test_reconcile_unaddressed_empty_when_all_addressed():
    accepted = [{'id': 'QH001', 'source': 'coherence', 'summary': 'x'}]
    edits = [{'source_item_id': 'QH001'}, {'source_item_id': 'QH001'}]  # two edits, same item
    assert reconcile_unaddressed(accepted, edits) == []


# ---------------------------------------------------------------------------
# Index + offset mapping
# ---------------------------------------------------------------------------

def test_build_index_and_para_index_at():
    paras_text = ["First line.", "Second line.", "Third line."]

    class P:
        def __init__(self, t): self.text = t
    paras = [P(t) for t in paras_text]
    texts, starts, full = build_index(paras)

    assert full == "First line.\nSecond line.\nThird line."
    assert starts == [0, 12, 25]
    # offset inside paragraph 1 ("Second line." starts at 12)
    assert _para_index_at(12, starts, texts) == 1
    assert _para_index_at(20, starts, texts) == 1
    # offset on the "\n" separator between para 0 and 1 -> None
    assert _para_index_at(11, starts, texts) is None


# ---------------------------------------------------------------------------
# Anchor location
# ---------------------------------------------------------------------------

def test_locate_anchor_exact_unique():
    full = "Alpha sentence. Beta sentence. Gamma sentence."
    method, s, e = locate_anchor("Beta sentence.", "", full)
    assert method == "exact"
    assert full[s:e] == "Beta sentence."


def test_locate_anchor_context_disambiguates_duplicate():
    full = "chase for status updates.\nMid text.\nchase for status updates."
    # Anchor appears twice; context_before picks the second occurrence.
    method, s, e = locate_anchor("chase for status updates.", "Mid text.\n", full)
    assert method == "context"
    assert s == full.rfind("chase for status updates.")


def test_locate_anchor_ambiguous_without_usable_context():
    full = "repeat me. repeat me."
    method, s, e = locate_anchor("repeat me.", "", full)
    assert method == "ambiguous"
    assert s is None


def test_locate_anchor_fuzzy_near_miss():
    full = "the firm is structurally weaker. Four factors drive the compression."
    # Anchor paraphrases the tail ("weakened." vs "weaker. Four...").
    anchor = "the firm is structurally weakened."
    method, s, e = locate_anchor(anchor, "", full)
    assert method == "fuzzy_near_miss"


def test_locate_anchor_none_when_text_absent():
    full = "completely unrelated content here."
    method, s, e = locate_anchor("this text is nowhere to be found in the document", "", full)
    assert method == "none"
    assert s is None


# ---------------------------------------------------------------------------
# In-place application
# ---------------------------------------------------------------------------

def test_replace_within_single_run():
    doc = Document()
    p = doc.add_paragraph("The margin fell $385,720 last year.")
    assert replace_in_paragraph(p, "$385,720", "$390,000") is True
    assert p.text == "The margin fell $390,000 last year."


def test_replace_across_runs_preserves_surrounding_formatting():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("The margin fell ")
    r2 = p.add_run("$385,720")
    r2.bold = True
    r3 = p.add_run(" last year.")
    # Replace the bold middle run's content.
    assert replace_in_paragraph(p, "$385,720", "$390,000") is True
    assert p.text == "The margin fell $390,000 last year."
    # Surrounding runs keep their (non-bold) formatting; replaced run keeps bold.
    assert r1.bold in (None, False)
    assert r2.bold is True


def test_replace_returns_false_when_anchor_absent():
    doc = Document()
    p = doc.add_paragraph("Nothing to change here.")
    assert replace_in_paragraph(p, "not present", "x") is False
    assert p.text == "Nothing to change here."


def test_insert_paragraph_after_adds_following_paragraph():
    doc = Document()
    p1 = doc.add_paragraph("First.")
    doc.add_paragraph("Second.")
    insert_paragraph_after(p1, "Inserted.")
    texts = [p.text for p in doc.paragraphs]
    assert texts == ["First.", "Inserted.", "Second."]


def test_iter_paragraphs_includes_table_cells_in_order():
    doc = Document()
    doc.add_paragraph("Body para.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    texts = [p.text for p in iter_paragraphs(doc)]
    assert "Body para." in texts
    assert "Cell A" in texts
    assert "Cell B" in texts

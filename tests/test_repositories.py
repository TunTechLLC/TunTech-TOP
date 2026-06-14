import sqlite3
import pytest
import os
from pathlib import Path


@pytest.fixture(autouse=True)
def test_db(monkeypatch, tmp_path):
    """Set up and tear down test database for each test.
    Uses tmp_path to get a unique path per test — avoids Windows file lock issues."""
    db_path = tmp_path / "test.db"
    monkeypatch.setenv("TOP_DB_PATH", str(db_path))

    conn = sqlite3.connect(db_path)
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS Clients (
            client_id TEXT PRIMARY KEY,
            firm_name TEXT NOT NULL,
            firm_size INTEGER,
            service_model TEXT,
            notes TEXT,
            created_date TEXT NOT NULL,
            confirmed_revenue REAL
        );
        CREATE TABLE IF NOT EXISTS Engagements (
            engagement_id TEXT PRIMARY KEY,
            client_id TEXT NOT NULL,
            engagement_name TEXT NOT NULL,
            status TEXT,
            start_date TEXT NOT NULL,
            end_date TEXT,
            engagement_type TEXT,
            stated_problem TEXT,
            client_hypothesis TEXT,
            previously_tried TEXT,
            notes TEXT,
            created_date TEXT NOT NULL,
            FOREIGN KEY (client_id) REFERENCES Clients(client_id)
        );
        CREATE TABLE IF NOT EXISTS Signals (
            signal_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            interview_id TEXT,
            signal_name TEXT NOT NULL,
            domain TEXT,
            observed_value TEXT,
            normalized_band TEXT,
            signal_confidence TEXT,
            economic_relevance TEXT,
            source TEXT,
            notes TEXT,
            created_date TEXT NOT NULL,
            source_file TEXT,
            library_signal_id TEXT,
            valence TEXT
        );
        CREATE TABLE IF NOT EXISTS Patterns (
            pattern_id TEXT PRIMARY KEY,
            pattern_name TEXT,
            domain TEXT,
            trigger_signals TEXT,
            operational_impact TEXT,
            likely_root_cause TEXT,
            recommended_improvements TEXT,
            economic_model TEXT,
            economic_formula TEXT
        );
        CREATE TABLE IF NOT EXISTS EngagementPatterns (
            ep_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            pattern_id TEXT,
            confidence TEXT,
            economic_impact_est TEXT,
            accepted INTEGER DEFAULT 0,
            notes TEXT,
            created_date TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS OPDFindings (
            finding_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            pattern_id TEXT,
            finding_title TEXT NOT NULL,
            domain TEXT,
            confidence TEXT,
            operational_impact TEXT,
            economic_impact TEXT,
            root_cause TEXT,
            recommendation TEXT,
            priority TEXT,
            effort TEXT,
            opd_section INTEGER,
            created_date TEXT NOT NULL,
            evidence_summary TEXT,
            key_quotes TEXT,
            display_figure TEXT,
            display_label TEXT,
            figure_type TEXT,
            include_in_executive INTEGER DEFAULT 0,
            confirmed_figure REAL,
            derived_figure REAL,
            annual_drag_figure REAL,
            valence TEXT
        );
        CREATE TABLE IF NOT EXISTS RoadmapItems (
            item_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            finding_id TEXT,
            initiative_name TEXT NOT NULL,
            domain TEXT,
            phase TEXT,
            priority TEXT,
            effort TEXT,
            estimated_impact TEXT,
            owner TEXT,
            target_date TEXT,
            status TEXT,
            created_date TEXT NOT NULL,
            capability TEXT,
            addressing_finding_ids TEXT,
            depends_on TEXT
        );
        CREATE TABLE IF NOT EXISTS KnowledgePromotions (
            promotion_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            finding_id TEXT,
            pattern_id TEXT,
            promotion_type TEXT NOT NULL,
            description TEXT NOT NULL,
            applied_to TEXT,
            promotion_date TEXT NOT NULL,
            created_date TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS AgentRuns (
            run_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            agent_name TEXT NOT NULL,
            model_used TEXT,
            run_date TEXT,
            output_summary TEXT,
            output_full TEXT,
            output_doc_link TEXT,
            accepted INTEGER DEFAULT 0,
            created_date TEXT NOT NULL,
            consultant_correction TEXT
        );
        CREATE TABLE IF NOT EXISTS ProcessedFiles (
            file_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            file_name TEXT NOT NULL,
            file_hash TEXT NOT NULL,
            file_type TEXT NOT NULL,
            processed_date TEXT NOT NULL,
            signal_count INTEGER DEFAULT 0,
            status TEXT DEFAULT 'processed',
            UNIQUE(file_hash)
        );
        CREATE TABLE IF NOT EXISTS QACoverageItems (
            qa_coverage_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            source_file TEXT NOT NULL,
            who_said_it TEXT NOT NULL,
            what_was_said TEXT NOT NULL,
            location_in_source TEXT NOT NULL,
            appears_in_roadmap INTEGER NOT NULL,
            roadmap_location TEXT,
            tier INTEGER NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            created_date TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS QACoherenceItems (
            qa_coherence_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            issue TEXT NOT NULL,
            category TEXT NOT NULL,
            sections_involved TEXT NOT NULL,
            recommended_fix TEXT NOT NULL,
            tier INTEGER NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            created_date TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS QAEditorialItems (
            qa_editorial_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            issue TEXT NOT NULL,
            category TEXT NOT NULL,
            location TEXT NOT NULL,
            recommended_fix TEXT NOT NULL,
            standard_term TEXT,
            tier INTEGER NOT NULL,
            source TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'pending',
            created_date TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS QARevisionEdits (
            qa_revision_id TEXT PRIMARY KEY,
            engagement_id TEXT NOT NULL,
            edit_type TEXT NOT NULL,
            qa_source TEXT NOT NULL,
            source_item_id TEXT,
            anchor TEXT NOT NULL,
            context_before TEXT,
            new_text TEXT NOT NULL,
            reason TEXT,
            outcome TEXT NOT NULL,
            match_method TEXT,
            created_date TEXT NOT NULL
        );
    """)
    conn.commit()
    conn.close()

    import importlib
    import config
    importlib.reload(config)

    yield


# ---------------------------------------------------------------------------
# Original tests
# ---------------------------------------------------------------------------

def test_engagement_repository_create_and_retrieve():
    """Repository integration test — create an engagement and read it back."""
    from api.db.repositories.engagement import EngagementRepository

    repo = EngagementRepository()
    data = {
        'firm_name':         'Test Firm',
        'firm_size':         50,
        'service_model':     'IT Consulting',
        'stated_problem':    'Test problem',
        'client_hypothesis': 'Test hypothesis',
        'previously_tried':  'Nothing',
    }

    engagement_id = repo.create(data)

    assert engagement_id is not None
    assert engagement_id.startswith('E')

    result = repo.get_by_id(engagement_id)
    assert result is not None
    assert result['firm_name'] == 'Test Firm'
    assert result['status'] == 'Active'

    all_engagements = repo.get_all()
    ids = [e['engagement_id'] for e in all_engagements]
    assert engagement_id in ids


def test_finding_repository_atomic_transaction():
    """Transaction test — finding creation and pattern acceptance are atomic."""
    from api.db.repositories.base import BaseRepository
    from api.db.repositories.finding import FindingRepository

    base         = BaseRepository()
    finding_repo = FindingRepository()

    base._write(
        "INSERT INTO EngagementPatterns VALUES (?,?,?,?,NULL,0,?,?)",
        ('EP_TEST', 'E_TEST', 'P38', 'High', 'test notes', '2026-03-23')
    )

    pattern_before = base._query(
        "SELECT accepted FROM EngagementPatterns WHERE ep_id = ?",
        ('EP_TEST',)
    )
    assert pattern_before[0]['accepted'] == 0

    finding_data = {
        'finding_title':      'Test Finding',
        'domain':             'Consulting Economics',
        'confidence':         'High',
        'operational_impact': 'Test operational impact',
        'economic_impact':    'Test economic impact',
        'root_cause':         'Test root cause',
        'recommendation':     'Test recommendation',
        'priority':           'High',
        'effort':             'Medium',
    }

    finding_id = finding_repo.create('E_TEST', finding_data, ['EP_TEST'])
    assert finding_id is not None

    pattern_after = base._query(
        "SELECT accepted FROM EngagementPatterns WHERE ep_id = ?",
        ('EP_TEST',)
    )
    assert pattern_after[0]['accepted'] == 1

    findings = finding_repo.get_all('E_TEST')
    assert len(findings) == 1
    assert findings[0]['finding_title'] == 'Test Finding'


# ---------------------------------------------------------------------------
# Valence round-trip (Strengths & Value-Case Reframe — Track 1)
# ---------------------------------------------------------------------------

def test_signal_valence_round_trips_and_defaults_null():
    """A signal's valence persists; absent valence defaults to NULL (unset)."""
    from api.db.repositories.signal import SignalRepository

    signal_repo = SignalRepository()
    base_signal = {
        'engagement_id':    'E_VAL',
        'signal_name':      'Strong win-rate trend',
        'domain':           'Sales & Pipeline',
        'observed_value':   '38% close rate',
        'normalized_band':  'Above benchmark',
        'signal_confidence': 'High',
        'source':           'Interview',
    }
    signal_repo.create({**base_signal, 'valence': 'Strength'})
    signal_repo.create({**base_signal, 'signal_name': 'Unlabeled signal'})  # no valence

    signals = {s['signal_name']: s for s in signal_repo.get_for_engagement('E_VAL')}
    assert signals['Strong win-rate trend']['valence'] == 'Strength'
    assert signals['Unlabeled signal']['valence'] is None


def test_finding_valence_round_trips_and_defaults_null():
    """A finding's valence persists; absent valence defaults to NULL (≡ Negative)."""
    from api.db.repositories.finding import FindingRepository

    finding_repo = FindingRepository()
    base = {
        'finding_title':      'Preserve Finding',
        'domain':             'Sales & Pipeline',
        'confidence':         'High',
        'operational_impact': 'x',
        'economic_impact':    'x',
        'root_cause':         'x',
        'recommendation':     'x',
        'valence':            'Positive',
    }
    # Positive findings have no contributing patterns — empty list is valid at the repo layer.
    finding_repo.create('E_VAL', base, [])
    finding_repo.create('E_VAL', {**base, 'finding_title': 'Plain Finding', 'valence': None}, [])

    findings = {f['finding_title']: f for f in finding_repo.get_all('E_VAL')}
    assert findings['Preserve Finding']['valence'] == 'Positive'
    assert findings['Plain Finding']['valence'] is None


def test_strength_evidence_chain_builds_from_domain_signals():
    """A Preserve finding's evidence chain has no patterns and carries the domain's
    Strength signals with verbatim quotes parsed from the notes format."""
    from api.routers.findings import _build_strength_evidence_chain

    strength_signals = [
        {
            'signal_id':        'S101',
            'signal_name':      'Repeat-client retention',
            'signal_confidence': 'High',
            'source_file':      'Interview_CEO.txt',
            'notes':            "Quote: 'seven of our top ten have renewed three years running' "
                                "— Interpretation: strong retention.",
        },
    ]
    chain = _build_strength_evidence_chain(strength_signals)
    assert chain['patterns'] == []
    assert chain['signals_hidden'] == 0
    assert len(chain['signals']) == 1
    entry = chain['signals'][0]
    assert entry['signal_id'] == 'S101'
    assert entry['signal_name'] == 'Repeat-client retention'
    assert 'seven of our top ten have renewed three years running' in entry['quote']
    assert 'Interpretation' not in entry['quote']  # interpretation stripped off the quote


def test_evidence_summary_positive_valence_reads_strength_backed():
    """Positive valence produces a strength-backed summary, not a pattern summary."""
    from api.routers.findings import _compute_evidence_summary

    summary = _compute_evidence_summary(
        [], [], 'Customer Experience', {},
        valence='positive', strength_count=3,
    )
    assert summary == 'Backed by 3 strength signals in Customer Experience'
    # singular
    one = _compute_evidence_summary([], [], 'Sales & Pipeline', {},
                                    valence='positive', strength_count=1)
    assert one == 'Backed by 1 strength signal in Sales & Pipeline'


def test_section_map_byte_identical_without_strengths():
    """With no Positive/Dual findings the section map equals the base map exactly —
    a negative-only deliverable keeps today's section numbering."""
    from api.services.report_sections import compute_section_map, _SECTION_MAP
    assert compute_section_map(False) == _SECTION_MAP
    assert 'what_to_preserve' not in compute_section_map(False)


def test_section_map_shifts_when_what_to_preserve_renders():
    """When strengths exist, 'What to Preserve' takes section 10 and the roadmap block
    shifts down by one; sections before it are untouched."""
    from api.services.report_sections import compute_section_map
    sm = compute_section_map(True)
    assert sm['what_to_preserve'] == 10
    assert sm['roadmap'] == 11
    assert sm['priority_zero'] == '11.1'
    assert sm['risks'] == '11.8'
    assert sm['what_happens_next'] == 12
    # Sections before the insertion point do not move.
    assert sm['domain_analysis'] == 6
    assert sm['economic_impact'] == 8
    assert sm['future_state'] == 9


def test_domain_score_positive_finding_does_not_subtract():
    """A Positive (Preserve) finding must not lower a domain's maturity score; a Negative
    finding subtracts exactly as before (backward-compatible)."""
    from api.services.report_sections import _compute_domain_scores
    sigs = [{'domain': 'Delivery Operations'}]
    neg = _compute_domain_scores(
        sigs, [], [{'domain': 'Delivery Operations', 'priority': 'High', 'valence': 'Negative'}]
    )
    pos = _compute_domain_scores(
        sigs, [], [{'domain': 'Delivery Operations', 'priority': 'High', 'valence': 'Positive'}]
    )
    assert neg['Delivery Operations'] == 4.5   # 5.0 - 0.5 for a High negative finding
    assert pos['Delivery Operations'] == 5.0   # strength does not deduct
    # NULL valence behaves like Negative (today's behavior)
    null = _compute_domain_scores(
        sigs, [], [{'domain': 'Delivery Operations', 'priority': 'High'}]
    )
    assert null['Delivery Operations'] == 4.5


def _render_finding(title, valence, domain='Delivery Operations'):
    return {
        'domain': domain, 'finding_title': title, 'confidence': 'High',
        'priority': 'High', 'effort': 'Medium', 'operational_impact': 'op',
        'economic_impact': 'econ', 'root_cause': 'root', 'recommendation': 'rec',
        'valence': valence, 'evidence_summary': '', 'key_quotes': '',
    }


def test_findings_by_domain_excludes_strength_findings():
    """Domain Analysis renders Negative findings only; Positive/Dual go to What to Preserve."""
    from docx import Document
    from api.services.report_generator import ReportGeneratorService
    svc = ReportGeneratorService('E_TEST')
    doc = Document()
    svc._findings_by_domain(doc, [
        _render_finding('Chronic Overruns', 'Negative'),
        _render_finding('Strong Retention', 'Positive', 'Customer Experience'),
        _render_finding('Win Rate Under Strain', 'Dual', 'Sales & Pipeline'),
    ], narrative={})
    text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Chronic Overruns' in text
    assert 'Strong Retention' not in text
    assert 'Win Rate Under Strain' not in text


def test_what_to_preserve_renders_strengths_and_omits_when_negative_only():
    """_what_to_preserve renders Positive + Dual findings; renders nothing for negatives."""
    from docx import Document
    from api.services.report_generator import ReportGeneratorService
    svc = ReportGeneratorService('E_TEST')

    doc = Document()
    svc._what_to_preserve(doc, [
        _render_finding('Strong Retention', 'Positive', 'Customer Experience'),
        _render_finding('Win Rate Under Strain', 'Dual', 'Sales & Pipeline'),
        _render_finding('Chronic Overruns', 'Negative'),  # ignored here
    ])
    text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'What to Preserve' in text
    assert 'Strong Retention' in text
    assert 'Win Rate Under Strain' in text
    assert 'Chronic Overruns' not in text

    # Negative-only → section omitted entirely (byte-identical to today)
    doc2 = Document()
    before = len(doc2.paragraphs)
    svc._what_to_preserve(doc2, [_render_finding('Chronic Overruns', 'Negative')])
    assert 'What to Preserve' not in '\n'.join(p.text for p in doc2.paragraphs)
    assert len(doc2.paragraphs) == before  # nothing added


def test_domain_cap_reserves_slots_for_strengths():
    """A risk-dominated domain must not crowd out its Strength signals — the cap reserves
    up to 2 strength slots so positive evidence reaches candidate review."""
    from api.services.document_processor import _apply_domain_cap
    cands = (
        [{'domain': 'Delivery Operations', 'signal_name': f'risk{i}',
          'signal_confidence': 'High', 'valence': 'Risk'} for i in range(5)]
        + [{'domain': 'Delivery Operations', 'signal_name': f'str{i}',
            'signal_confidence': 'Medium', 'valence': 'Strength'} for i in range(2)]
    )
    capped, removed = _apply_domain_cap(cands, cap=5)
    assert len(capped) == 5                       # cap total preserved
    assert removed == 2
    assert sum(1 for c in capped if c['valence'] == 'Strength') == 2   # both strengths kept


def test_accept_enforces_single_accepted_run_per_agent():
    """Re-running and re-accepting an agent must leave exactly one accepted run — the
    latest — so get_accepted_output is deterministic. Regression for the bug where a
    re-run left two accepted runs and the truncated one could be read downstream."""
    from api.db.repositories.agent_run import AgentRunRepository
    repo = AgentRunRepository()

    first  = repo.create({'engagement_id': 'E_ACC', 'agent_name': 'Synthesizer',
                          'output_full': 'FIRST run output (truncated)'})
    second = repo.create({'engagement_id': 'E_ACC', 'agent_name': 'Synthesizer',
                          'output_full': 'SECOND run output (complete)'})

    repo.accept(first)
    repo.accept(second)   # should un-accept `first`

    runs = [r for r in repo.get_for_engagement('E_ACC') if r['agent_name'] == 'Synthesizer']
    accepted = [r for r in runs if r['accepted'] == 1]
    assert len(accepted) == 1, f"expected exactly one accepted run, got {len(accepted)}"
    assert accepted[0]['run_id'] == second
    assert repo.get_accepted_output('E_ACC', 'Synthesizer') == 'SECOND run output (complete)'


def test_domain_cap_unchanged_without_strengths():
    """With no Strength signals the cap behaves exactly as before — top N by confidence."""
    from api.services.document_processor import _apply_domain_cap
    cands = [{'domain': 'Sales & Pipeline', 'signal_name': f's{i}',
              'signal_confidence': 'High', 'valence': 'Risk'} for i in range(7)]
    capped, removed = _apply_domain_cap(cands, cap=5)
    assert len(capped) == 5
    assert removed == 2
    assert all(c['valence'] == 'Risk' for c in capped)


# ---------------------------------------------------------------------------
# Phase 8 tests
# ---------------------------------------------------------------------------

def test_pattern_bulk_create_sequential_ids():
    """Verify bulk_create generates unique sequential EP IDs.
    Regression test for the duplicate ID bug — was using list comprehension
    which called next_ep_id() before any rows were written."""
    from api.db.repositories.engagement import EngagementRepository
    from api.db.repositories.pattern import PatternRepository
    from api.db.repositories.base import BaseRepository

    # Insert a test pattern into the library so the FK is valid
    base = BaseRepository()
    base._write(
        "INSERT INTO Patterns (pattern_id, pattern_name, domain) VALUES (?,?,?)",
        ('P12', 'Test Pattern 12', 'Delivery Operations')
    )
    base._write(
        "INSERT INTO Patterns (pattern_id, pattern_name, domain) VALUES (?,?,?)",
        ('P27', 'Test Pattern 27', 'Consulting Economics')
    )
    base._write(
        "INSERT INTO Patterns (pattern_id, pattern_name, domain) VALUES (?,?,?)",
        ('P38', 'Test Pattern 38', 'Resource Management')
    )

    # Create an engagement to use as FK
    eng_repo = EngagementRepository()
    engagement_id = eng_repo.create({
        'firm_name':         'Test Firm',
        'firm_size':         50,
        'service_model':     'IT Consulting',
        'stated_problem':    'Test',
        'client_hypothesis': '',
        'previously_tried':  '',
    })

    pattern_repo = PatternRepository()
    rows = [
        {'engagement_id': engagement_id, 'pattern_id': 'P12', 'confidence': 'High',   'notes': 'test1'},
        {'engagement_id': engagement_id, 'pattern_id': 'P27', 'confidence': 'Medium', 'notes': 'test2'},
        {'engagement_id': engagement_id, 'pattern_id': 'P38', 'confidence': 'High',   'notes': 'test3'},
    ]
    pattern_repo.bulk_create(rows)

    patterns = pattern_repo.get_for_engagement(engagement_id)
    ep_ids   = [p['ep_id'] for p in patterns]

    assert len(ep_ids) == 3, f"Expected 3 patterns, got {len(ep_ids)}"
    assert len(ep_ids) == len(set(ep_ids)), f"Duplicate EP IDs generated: {ep_ids}"


def test_agent_run_prerequisites_blocking():
    """Verify validate_prerequisites blocks agents when required agents not accepted."""
    from api.db.repositories.engagement import EngagementRepository
    from api.db.repositories.agent_run import AgentRunRepository

    eng_repo = EngagementRepository()
    engagement_id = eng_repo.create({
        'firm_name':         'Test Firm',
        'firm_size':         50,
        'service_model':     'IT Consulting',
        'stated_problem':    'Test',
        'client_hypothesis': '',
        'previously_tried':  '',
    })

    agent_repo = AgentRunRepository()

    # No agents run yet — all three prerequisites should be missing
    missing = agent_repo.validate_prerequisites(
        engagement_id,
        ['Diagnostician', 'Delivery Operations', 'Consulting Economics']
    )
    assert 'Diagnostician' in missing
    assert 'Delivery Operations' in missing
    assert 'Consulting Economics' in missing

    # Create and accept a Diagnostician run
    run_id = agent_repo.create({
        'engagement_id':  engagement_id,
        'agent_name':     'Diagnostician',
        'output_full':    'Full test output for Diagnostician',
        'output_summary': 'Test summary',
        'model_used':     'claude-sonnet-4-6',
    })
    agent_repo.accept(run_id)

    # Now only Delivery Operations and Consulting Economics should be missing
    missing = agent_repo.validate_prerequisites(
        engagement_id,
        ['Diagnostician', 'Delivery Operations', 'Consulting Economics']
    )
    assert 'Diagnostician' not in missing
    assert 'Delivery Operations' in missing
    assert 'Consulting Economics' in missing


def test_pattern_detection_result_validators():
    """Verify Pydantic validators reject invalid pattern_ids and confidence values."""
    from api.models.pattern import PatternDetectionResult

    # Valid data should pass
    valid = PatternDetectionResult(pattern_id='P12', confidence='High', notes='test')
    assert valid.pattern_id == 'P12'
    assert valid.confidence == 'High'

    # Valid hypothesis confidence should pass
    valid2 = PatternDetectionResult(pattern_id='P5', confidence='Hypothesis', notes='weak signal')
    assert valid2.confidence == 'Hypothesis'

    # Invalid pattern_id format should raise
    with pytest.raises(Exception):
        PatternDetectionResult(pattern_id='INVALID', confidence='High', notes='test')

    # Invalid confidence value should raise
    with pytest.raises(Exception):
        PatternDetectionResult(pattern_id='P12', confidence='Excellent', notes='test')

    # Non-P prefix should raise
    with pytest.raises(Exception):
        PatternDetectionResult(pattern_id='EP012', confidence='High', notes='test')


# ---------------------------------------------------------------------------
# QA-1 Coverage Check tests
# ---------------------------------------------------------------------------

def _make_qa_test_engagement():
    """Helper — create a test engagement for QA coverage tests."""
    from api.db.repositories.engagement import EngagementRepository
    return EngagementRepository().create({
        'firm_name':         'QA Test Firm',
        'firm_size':         50,
        'service_model':     'IT Consulting',
        'stated_problem':    'Test',
        'client_hypothesis': '',
        'previously_tried':  '',
    })


def _sample_item(tier: int, appears: int = 0, roadmap_location=None):
    """Helper — minimal valid QA coverage item dict for tests."""
    item = {
        'source_file':        'Interview_CEO.txt',
        'who_said_it':        'CEO',
        'what_was_said':      f'Tier {tier} test item',
        'location_in_source': 'Section A; Lines 1-5',
        'appears_in_roadmap': appears,
        'tier':               tier,
    }
    if roadmap_location is not None:
        item['roadmap_location'] = roadmap_location
    return item


def test_qa_coverage_bulk_create_sequential_ids():
    """Verify bulk_create generates unique sequential QC IDs.
    Same regression pattern as test_pattern_bulk_create_sequential_ids —
    sequential loop required to avoid duplicate IDs."""
    from api.db.repositories.qa_coverage import QACoverageRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoverageRepository()

    items = [
        _sample_item(1),
        _sample_item(2),
        _sample_item(3),
        _sample_item(1, appears=1, roadmap_location='Section 4.2'),
    ]
    count = repo.bulk_create(engagement_id, items)

    rows = repo.get_for_engagement(engagement_id)
    qc_ids = [r['qa_coverage_id'] for r in rows]

    assert count == 4
    assert len(qc_ids) == 4
    assert len(qc_ids) == len(set(qc_ids)), f"Duplicate QC IDs generated: {qc_ids}"
    for qc_id in qc_ids:
        assert qc_id.startswith('QC')
        assert qc_id[2:].isdigit()


def test_qa_coverage_get_for_engagement_ordered_by_tier():
    """Items should be returned ordered by tier ASC, then qa_coverage_id ASC."""
    from api.db.repositories.qa_coverage import QACoverageRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoverageRepository()

    # Deliberately insert in non-tier order
    repo.bulk_create(engagement_id, [
        _sample_item(3),
        _sample_item(1),
        _sample_item(2),
        _sample_item(1),
    ])

    rows = repo.get_for_engagement(engagement_id)
    tiers = [r['tier'] for r in rows]
    assert tiers == [1, 1, 2, 3], f"Expected tier order [1,1,2,3], got {tiers}"


def test_qa_coverage_update_status():
    """Single-item status update writes through and is visible on next read."""
    from api.db.repositories.qa_coverage import QACoverageRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoverageRepository()
    repo.bulk_create(engagement_id, [_sample_item(2)])

    [item] = repo.get_for_engagement(engagement_id)
    assert item['status'] == 'pending'

    repo.update_status(item['qa_coverage_id'], 'accepted')
    [item_after] = repo.get_for_engagement(engagement_id)
    assert item_after['status'] == 'accepted'


def test_qa_coverage_batch_accept_tier_1_respects_explicit_rejection():
    """batch_accept_tier_1 must only move pending Tier 1 items to accepted.
    Rejected items stay rejected. Tier 2 and Tier 3 items are untouched."""
    from api.db.repositories.qa_coverage import QACoverageRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoverageRepository()
    repo.bulk_create(engagement_id, [
        _sample_item(1),  # pending Tier 1 → should become accepted
        _sample_item(1),  # will be rejected explicitly → must stay rejected
        _sample_item(2),  # pending Tier 2 → must stay pending
        _sample_item(3),  # pending Tier 3 → must stay pending
    ])

    items = repo.get_for_engagement(engagement_id)
    tier1_to_reject = next(i for i in items if i['tier'] == 1 and i == items[1])
    repo.update_status(tier1_to_reject['qa_coverage_id'], 'rejected')

    updated = repo.batch_accept_tier_1(engagement_id)
    assert updated == 1, f"Expected 1 row updated (the pending Tier 1), got {updated}"

    items_after = repo.get_for_engagement(engagement_id)
    by_id = {i['qa_coverage_id']: i for i in items_after}

    accepted_t1 = [i for i in items_after if i['tier'] == 1 and i['status'] == 'accepted']
    rejected_t1 = [i for i in items_after if i['tier'] == 1 and i['status'] == 'rejected']
    pending_t2  = [i for i in items_after if i['tier'] == 2 and i['status'] == 'pending']
    pending_t3  = [i for i in items_after if i['tier'] == 3 and i['status'] == 'pending']

    assert len(accepted_t1) == 1
    assert len(rejected_t1) == 1
    assert len(pending_t2)  == 1
    assert len(pending_t3)  == 1
    assert by_id[tier1_to_reject['qa_coverage_id']]['status'] == 'rejected'


def test_qa_coverage_delete_for_engagement_clears_prior_run():
    """delete_for_engagement is called before re-detection — must remove all items
    for the engagement so the new run starts from a clean slate."""
    from api.db.repositories.qa_coverage import QACoverageRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoverageRepository()
    repo.bulk_create(engagement_id, [
        _sample_item(1),
        _sample_item(2),
        _sample_item(3),
    ])
    assert len(repo.get_for_engagement(engagement_id)) == 3

    deleted = repo.delete_for_engagement(engagement_id)
    assert deleted == 3
    assert repo.get_for_engagement(engagement_id) == []


# ---------------------------------------------------------------------------
# QA-2 Coherence Check tests
# ---------------------------------------------------------------------------

def _coherence_sample(tier: int, category: str = 'contradiction',
                       sections=None):
    """Helper — minimal valid coherence item dict for tests."""
    return {
        'issue':             f'Tier {tier} {category} test issue',
        'category':          category,
        'sections_involved': sections if sections is not None else ['Section A', 'Section B'],
        'recommended_fix':   f'Apply the fix for {category}',
        'tier':              tier,
    }


def test_qa_coherence_bulk_create_sequential_ids():
    """Verify bulk_create generates unique sequential QH IDs.
    Same regression pattern as the pattern/coverage bulk_create tests —
    sequential loop required for unique IDs."""
    from api.db.repositories.qa_coherence import QACoherenceRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoherenceRepository()
    count = repo.bulk_create(engagement_id, [
        _coherence_sample(1, 'contradiction'),
        _coherence_sample(2, 'priority_mismatch'),
        _coherence_sample(3, 'weak_grounding'),
        _coherence_sample(1, 'missing_root_cause'),
    ])

    rows = repo.get_for_engagement(engagement_id)
    qh_ids = [r['qa_coherence_id'] for r in rows]

    assert count == 4
    assert len(qh_ids) == 4
    assert len(qh_ids) == len(set(qh_ids)), f"Duplicate QH IDs generated: {qh_ids}"
    for qh_id in qh_ids:
        assert qh_id.startswith('QH')
        assert qh_id[2:].isdigit()


def test_qa_coherence_get_orders_by_tier_and_deserializes_sections():
    """Items are returned in tier order; sections_involved is parsed back to a list."""
    from api.db.repositories.qa_coherence import QACoherenceRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoherenceRepository()
    repo.bulk_create(engagement_id, [
        _coherence_sample(3, sections=['Table 19', 'Table 25']),
        _coherence_sample(1, sections=['Table 3', 'Resource Management summary']),
        _coherence_sample(2, sections=['Table 11', 'Table 25']),
    ])

    rows = repo.get_for_engagement(engagement_id)
    assert [r['tier'] for r in rows] == [1, 2, 3]
    assert rows[0]['sections_involved'] == ['Table 3', 'Resource Management summary']
    assert isinstance(rows[0]['sections_involved'], list)


def test_qa_coherence_update_status():
    """Single-item status update writes through."""
    from api.db.repositories.qa_coherence import QACoherenceRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoherenceRepository()
    repo.bulk_create(engagement_id, [_coherence_sample(2)])

    [item] = repo.get_for_engagement(engagement_id)
    assert item['status'] == 'pending'

    repo.update_status(item['qa_coherence_id'], 'accepted')
    [after] = repo.get_for_engagement(engagement_id)
    assert after['status'] == 'accepted'


def test_qa_coherence_batch_accept_tier_1_respects_rejection():
    """batch_accept_tier_1 only moves pending Tier 1 items to accepted —
    rejected items stay rejected; Tier 2/3 untouched."""
    from api.db.repositories.qa_coherence import QACoherenceRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoherenceRepository()
    repo.bulk_create(engagement_id, [
        _coherence_sample(1, 'contradiction'),       # will stay pending then become accepted
        _coherence_sample(1, 'priority_mismatch'),   # will be rejected explicitly
        _coherence_sample(2, 'weak_grounding'),
        _coherence_sample(3, 'missing_root_cause'),
    ])

    items = repo.get_for_engagement(engagement_id)
    tier1_to_reject = [i for i in items if i['tier'] == 1][1]
    repo.update_status(tier1_to_reject['qa_coherence_id'], 'rejected')

    updated = repo.batch_accept_tier_1(engagement_id)
    assert updated == 1

    after = repo.get_for_engagement(engagement_id)
    accepted_t1 = [i for i in after if i['tier'] == 1 and i['status'] == 'accepted']
    rejected_t1 = [i for i in after if i['tier'] == 1 and i['status'] == 'rejected']
    pending_t2  = [i for i in after if i['tier'] == 2 and i['status'] == 'pending']
    pending_t3  = [i for i in after if i['tier'] == 3 and i['status'] == 'pending']
    assert len(accepted_t1) == 1
    assert len(rejected_t1) == 1
    assert len(pending_t2)  == 1
    assert len(pending_t3)  == 1


def test_qa_coherence_delete_for_engagement_clears_prior_run():
    """delete_for_engagement removes all items for the engagement."""
    from api.db.repositories.qa_coherence import QACoherenceRepository

    engagement_id = _make_qa_test_engagement()
    repo = QACoherenceRepository()
    repo.bulk_create(engagement_id, [
        _coherence_sample(1),
        _coherence_sample(2),
        _coherence_sample(3),
    ])
    assert len(repo.get_for_engagement(engagement_id)) == 3

    deleted = repo.delete_for_engagement(engagement_id)
    assert deleted == 3
    assert repo.get_for_engagement(engagement_id) == []


# ---------------------------------------------------------------------------
# QA-3 Editorial Check tests
# ---------------------------------------------------------------------------

def _editorial_sample(tier: int, source: str = 'python', category: str = 'context_gap'):
    """Helper — minimal valid editorial item dict for tests."""
    return {
        'issue':           f'Tier {tier} {source} {category} test issue',
        'category':        category,
        'location':        f'Test location for tier {tier}',
        'recommended_fix': f'Apply fix for tier {tier} {source} item',
        'standard_term':   None,
        'tier':            tier,
        'source':          source,
    }


def test_qa_editorial_bulk_create_sequential_ids():
    """Verify bulk_create generates unique sequential QE IDs."""
    from api.db.repositories.qa_editorial import QAEditorialRepository

    engagement_id = _make_qa_test_engagement()
    repo = QAEditorialRepository()
    count = repo.bulk_create(engagement_id, [
        _editorial_sample(1, 'python', 'context_gap'),
        _editorial_sample(2, 'claude', 'voice'),
        _editorial_sample(3, 'python', 'terminology'),
        _editorial_sample(1, 'claude', 'voice'),
    ])

    rows = repo.get_for_engagement(engagement_id)
    qe_ids = [r['qa_editorial_id'] for r in rows]
    assert count == 4
    assert len(qe_ids) == len(set(qe_ids)), f"Duplicate QE IDs: {qe_ids}"
    for qe_id in qe_ids:
        assert qe_id.startswith('QE')
        assert qe_id[2:].isdigit()


def test_qa_editorial_get_orders_by_tier_and_preserves_source():
    """Items returned in tier order; source field roundtrips correctly."""
    from api.db.repositories.qa_editorial import QAEditorialRepository

    engagement_id = _make_qa_test_engagement()
    repo = QAEditorialRepository()
    repo.bulk_create(engagement_id, [
        _editorial_sample(3, 'claude'),
        _editorial_sample(1, 'python'),
        _editorial_sample(2, 'claude'),
    ])

    rows = repo.get_for_engagement(engagement_id)
    assert [r['tier'] for r in rows] == [1, 2, 3]
    assert rows[0]['source'] == 'python'
    assert rows[1]['source'] == 'claude'


def test_qa_editorial_update_status():
    """Single-item status update writes through."""
    from api.db.repositories.qa_editorial import QAEditorialRepository

    engagement_id = _make_qa_test_engagement()
    repo = QAEditorialRepository()
    repo.bulk_create(engagement_id, [_editorial_sample(2)])

    [item] = repo.get_for_engagement(engagement_id)
    assert item['status'] == 'pending'

    repo.update_status(item['qa_editorial_id'], 'accepted')
    [after] = repo.get_for_engagement(engagement_id)
    assert after['status'] == 'accepted'


def test_qa_editorial_batch_accept_tier_1_respects_rejection():
    """batch_accept_tier_1 only moves pending Tier 1 to accepted; rejected stays."""
    from api.db.repositories.qa_editorial import QAEditorialRepository

    engagement_id = _make_qa_test_engagement()
    repo = QAEditorialRepository()
    repo.bulk_create(engagement_id, [
        _editorial_sample(1, 'python'),
        _editorial_sample(1, 'claude'),
        _editorial_sample(2, 'python'),
        _editorial_sample(3, 'claude'),
    ])

    items = repo.get_for_engagement(engagement_id)
    tier1_to_reject = [i for i in items if i['tier'] == 1][1]
    repo.update_status(tier1_to_reject['qa_editorial_id'], 'rejected')

    updated = repo.batch_accept_tier_1(engagement_id)
    assert updated == 1

    after = repo.get_for_engagement(engagement_id)
    accepted_t1 = [i for i in after if i['tier'] == 1 and i['status'] == 'accepted']
    rejected_t1 = [i for i in after if i['tier'] == 1 and i['status'] == 'rejected']
    assert len(accepted_t1) == 1
    assert len(rejected_t1) == 1


def test_qa_editorial_delete_for_engagement_clears_prior_run():
    """delete_for_engagement removes all items for the engagement."""
    from api.db.repositories.qa_editorial import QAEditorialRepository

    engagement_id = _make_qa_test_engagement()
    repo = QAEditorialRepository()
    repo.bulk_create(engagement_id, [
        _editorial_sample(1),
        _editorial_sample(2),
        _editorial_sample(3),
    ])
    assert len(repo.get_for_engagement(engagement_id)) == 3

    deleted = repo.delete_for_engagement(engagement_id)
    assert deleted == 3
    assert repo.get_for_engagement(engagement_id) == []


# ---------------------------------------------------------------------------
# QA-3 Python editorial check function tests
# ---------------------------------------------------------------------------

def test_editorial_signal_codes_in_prose_catches_codes_and_dedupes():
    """check_signal_codes_in_prose flags S\\d{3,4} occurrences. Same code
    appearing multiple times produces one item (deduplication)."""
    from api.services.editorial_auditor import check_signal_codes_in_prose

    text = (
        "The Customer Experience finding is CONFIRMED per S217. "
        "Pricing Governance is CONFIRMED per Director Observation #4 and S237. "
        "Resource Management cites S217 again later in the same finding card."
    )
    items = check_signal_codes_in_prose(text)
    assert len(items) == 2, f"Expected 2 unique codes (S217, S237), got {len(items)}"
    codes_flagged = {item['issue'].split("'")[1] for item in items}
    assert codes_flagged == {'S217', 'S237'}
    for item in items:
        assert item['tier'] == 1
        assert item['source'] == 'python'
        assert item['category'] == 'context_gap'


def test_editorial_signal_codes_skips_project_ids_like_P10():
    """The check only matches S-codes (3-4 digit), not P-codes which collide
    with client project IDs (P10, P14, etc. in source materials)."""
    from api.services.editorial_auditor import check_signal_codes_in_prose

    text = (
        "P14 Blue Sky Supply Chain requires a PM assignment this week. "
        "The S217 reference is internal. "
        "Project P10 Cascade Health is also at risk."
    )
    items = check_signal_codes_in_prose(text)
    assert len(items) == 1
    assert "'S217'" in items[0]['issue']


def test_editorial_undefined_acronym_flags_PMO_without_definition():
    """check_undefined_acronyms_at_first_use flags PMO when 'Project
    Management Office' is missing or appears later."""
    from api.services.editorial_auditor import check_undefined_acronyms_at_first_use

    # Case 1: full form never appears
    text1 = "The PMO governance framework must be established with PM accountability built in."
    items1 = check_undefined_acronyms_at_first_use(text1)
    pmo_items = [i for i in items1 if "'PMO'" in i['issue']]
    assert len(pmo_items) == 1
    assert pmo_items[0]['standard_term'] == 'Project Management Office'
    assert pmo_items[0]['tier'] == 1

    # Case 2: full form appears AFTER abbreviation
    text2 = (
        "The PMO governance framework must be established. "
        "The Project Management Office will own this responsibility."
    )
    items2 = check_undefined_acronyms_at_first_use(text2)
    pmo_items_2 = [i for i in items2 if "'PMO'" in i['issue']]
    assert len(pmo_items_2) == 1  # still flagged — definition comes too late

    # Case 3: properly introduced — full form FIRST
    text3 = (
        "The Project Management Office (PMO) will own delivery standards. "
        "Subsequent PMO functions include resource allocation."
    )
    items3 = check_undefined_acronyms_at_first_use(text3)
    pmo_items_3 = [i for i in items3 if "'PMO'" in i['issue']]
    assert len(pmo_items_3) == 0


def test_editorial_operations_role_drift_requires_both_terms():
    """check_operations_role_drift only flags when BOTH 'Operations Manager'
    and 'Director of Operations' appear in the same document."""
    from api.services.editorial_auditor import check_operations_role_drift

    # Both terms — should flag
    text_both = (
        "The Operations Manager is responsible for tooling. "
        "Sandra Okafor, Director of Operations, owns the policy."
    )
    items_both = check_operations_role_drift(text_both)
    assert len(items_both) == 1
    assert items_both[0]['standard_term'] == 'Director of Operations'
    assert items_both[0]['tier'] == 1

    # Only one term — should NOT flag
    text_one = "Sandra Okafor, Director of Operations, owns the policy."
    items_one = check_operations_role_drift(text_one)
    assert items_one == []


# ---------------------------------------------------------------------------
# QA-4 Revision Agent — repository tests
# ---------------------------------------------------------------------------

def _revision_edit(outcome='applied', edit_type='replace', qa_source='coherence',
                   match_method='exact'):
    """Helper — minimal valid QA revision edit dict for tests."""
    return {
        'edit_type':      edit_type,
        'qa_source':      qa_source,
        'source_item_id': 'QH004',
        'anchor':         'Three actions must happen this week.',
        'context_before': 'Priority Zero. ',
        'new_text':       'Four actions must happen this week.',
        'reason':         'Coherence QH004: Priority Zero is four actions.',
        'outcome':        outcome,
        'match_method':   match_method,
    }


def test_qa_revision_bulk_create_sequential_ids():
    """bulk_create generates unique sequential QR IDs (no duplicate-ID regression)."""
    from api.db.repositories.qa_revision import QARevisionRepository

    engagement_id = _make_qa_test_engagement()
    repo = QARevisionRepository()
    count = repo.bulk_create(engagement_id, [
        _revision_edit('applied'),
        _revision_edit('flagged_unresolved', match_method='none'),
        _revision_edit('manual', edit_type='manual', match_method=None),
    ])

    rows = repo.get_for_engagement(engagement_id)
    qr_ids = [r['qa_revision_id'] for r in rows]
    assert count == 3
    assert len(qr_ids) == len(set(qr_ids)), f"Duplicate QR IDs: {qr_ids}"
    for qr_id in qr_ids:
        assert qr_id.startswith('QR')
        assert qr_id[2:].isdigit()


def test_qa_revision_persists_outcome_and_match_method():
    """Outcome and match_method round-trip through storage."""
    from api.db.repositories.qa_revision import QARevisionRepository

    engagement_id = _make_qa_test_engagement()
    repo = QARevisionRepository()
    repo.bulk_create(engagement_id, [
        _revision_edit('applied', match_method='context'),
        _revision_edit('flagged_unresolved', match_method='fuzzy_near_miss'),
    ])

    rows = repo.get_for_engagement(engagement_id)
    outcomes = {r['outcome'] for r in rows}
    methods = {r['match_method'] for r in rows}
    assert outcomes == {'applied', 'flagged_unresolved'}
    assert methods == {'context', 'fuzzy_near_miss'}
    assert all(r['source_item_id'] == 'QH004' for r in rows)


def test_qa_revision_update_outcome():
    """update_outcome lets QA-5 mark a flagged item as manually handled."""
    from api.db.repositories.qa_revision import QARevisionRepository

    engagement_id = _make_qa_test_engagement()
    repo = QARevisionRepository()
    repo.bulk_create(engagement_id, [_revision_edit('flagged_unresolved', match_method='none')])

    [edit] = repo.get_for_engagement(engagement_id)
    repo.update_outcome(edit['qa_revision_id'], 'manual_done')
    [after] = repo.get_for_engagement(engagement_id)
    assert after['outcome'] == 'manual_done'


def test_qa_revision_delete_for_engagement_clears_prior_run():
    """A re-run replaces the prior revision record."""
    from api.db.repositories.qa_revision import QARevisionRepository

    engagement_id = _make_qa_test_engagement()
    repo = QARevisionRepository()
    repo.bulk_create(engagement_id, [_revision_edit(), _revision_edit()])
    assert len(repo.get_for_engagement(engagement_id)) == 2

    deleted = repo.delete_for_engagement(engagement_id)
    assert deleted == 2
    assert repo.get_for_engagement(engagement_id) == []


# --- Defect A2: QA JSON parser raises on failure, never masks it as an empty result ---

def test_parse_json_array_empty_array_is_valid_not_failure():
    """A legitimately-empty model result '[]' parses to [] — distinct from a failure."""
    from api.services.claude import _parse_json_array
    assert _parse_json_array("[]", "ctx") == []


def test_parse_json_array_parses_valid_and_fenced():
    from api.services.claude import _parse_json_array
    assert _parse_json_array('[{"a": 1}]', "ctx") == [{"a": 1}]
    assert _parse_json_array('```[{"a": 1}]```', "ctx") == [{"a": 1}]  # fence stripped


def test_parse_json_array_raises_on_no_array():
    """No JSON array present (e.g. model returned prose) RAISES — never returns []."""
    import pytest
    from api.services.claude import _parse_json_array
    with pytest.raises(ValueError):
        _parse_json_array("there were no issues found", "ctx")


def test_parse_json_array_raises_on_malformed():
    """Malformed JSON inside the brackets RAISES — a failed QA check can't read as clean."""
    import json
    import pytest
    from api.services.claude import _parse_json_array
    with pytest.raises(json.JSONDecodeError):
        _parse_json_array('[ {"a":1} {"b":2} ]', "ctx")  # missing comma

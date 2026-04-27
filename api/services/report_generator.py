import os
import logging
import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from api.db.repositories.engagement import EngagementRepository
from api.db.repositories.finding import FindingRepository
from api.db.repositories.pattern import PatternRepository
from api.db.repositories.roadmap import RoadmapRepository
from api.db.repositories.agent_run import AgentRunRepository
from api.db.repositories.processed_files import ProcessedFilesRepository
from api.db.repositories.reporting import ReportingRepository
from api.services.claude import generate_report_narrative
from api.services.report_sections import (
    ReportSectionsMixin,
    _resolve_initiative_codes,
    _extract_interview_roles,
    _extract_document_types,
    _parse_display_figure_to_float,
    _prepopulate_display_figure,   # re-exported — api/routers/findings.py imports from here
    _SECTION_MAP,
)

logger = logging.getLogger(__name__)

_TEMPLATE = os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'roadmap_template.docx')


class ReportGeneratorService(ReportSectionsMixin):
    """Generates the OPD Transformation Roadmap Word document.

    Nine sections (three-layer design — CEO / Leadership / Execution):
    1. Executive Summary — opening para + Key Findings box + 3 short paras + reference line
    How to Read This Document — prefatory page, excluded from TOC
    2. Engagement Overview — compact metadata + narrator paragraph
    3. Operational Maturity Overview — intro paragraph + renamed signal table + callout
    4. Domain Analysis — role callout + narrator opening/closing + finding tables per domain
    5. Root Cause Analysis — narrator prose only
    6. Economic Impact Analysis — economic summary table + narrator prose
    7. Future State — metrics table (narrator) + narrative (narrator)
    8. Transformation Roadmap — 8.1 Priority Zero | 8.2 Overview | 8.3 Quick Wins (conditional)
                                 | 8.4-8.6 Phase Tables | 8.7 Dependencies | 8.8 Key Risks
    9. Immediate Next Steps — action table (narrator, execution-voice completion criteria)
    """

    def __init__(self, engagement_id: str):
        self.engagement_id = engagement_id

    async def generate(self) -> str:
        """Generate the OPD Word document. Returns the saved file path.
        Raises ValueError if the engagement is not found or has no accepted Synthesizer output."""
        eng = EngagementRepository().get_by_id(self.engagement_id)
        if not eng:
            raise ValueError(f"Engagement {self.engagement_id} not found")

        synth_output = AgentRunRepository().get_accepted_output(self.engagement_id, "Synthesizer")
        if not synth_output:
            raise ValueError(
                f"Engagement {self.engagement_id} has no accepted Synthesizer output. "
                "Complete and accept all five agents before generating the report."
            )

        findings        = FindingRepository().get_all(self.engagement_id)
        roadmap         = RoadmapRepository().get_all(self.engagement_id)
        signals         = ReportingRepository().get_engagement_signals(self.engagement_id)
        patterns        = PatternRepository().get_for_engagement(self.engagement_id)
        processed_files = ProcessedFilesRepository().get_for_engagement(self.engagement_id)

        interview_roles = _extract_interview_roles(processed_files)
        document_types  = _extract_document_types(processed_files)
        total_signals   = len(signals)
        domain_count    = len({s['domain'] for s in signals if s.get('domain')})

        # Build cross-reference strings from _SECTION_MAP — single source of truth.
        # Passed to the narrator so the prompt never hardcodes section numbers.
        section_refs = {
            'domain_analysis_ref': (
                f"(see Section {_SECTION_MAP['domain_analysis']} — Domain Analysis for full findings)"
            ),
            'economic_impact_ref': (
                f"(see Section {_SECTION_MAP['economic_impact']} — Economic Impact Analysis)"
            ),
            'priority_zero_ref': (
                f"(see Section {_SECTION_MAP['priority_zero']} — Priority Zero Actions)"
            ),
        }

        narrative = await generate_report_narrative(
            synth_output, findings, roadmap, eng,
            interview_roles=interview_roles,
            document_types=document_types,
            total_signals=total_signals,
            domain_count=domain_count,
            section_refs=section_refs,
        )

        if os.path.exists(_TEMPLATE):
            doc = Document(_TEMPLATE)
        else:
            logger.warning(f"Template not found at {_TEMPLATE} — using default styles")
            doc = Document()
        self._build(doc, eng, findings, roadmap, signals, patterns, narrative)

        file_path = self._output_path(eng)
        doc.save(file_path)
        logger.info(f"Report saved: {file_path}")
        return file_path

    # ------------------------------------------------------------------
    # File path helpers
    # ------------------------------------------------------------------

    def _output_path(self, eng: dict) -> str:
        """Save report to the engagement's reports_folder.
        Falls back to the system temp directory if reports_folder is not set."""
        reports_dir = eng.get('reports_folder') or ''
        if reports_dir:
            os.makedirs(reports_dir, exist_ok=True)
            return os.path.join(reports_dir, f"OPD_Transformation_Roadmap_{self.engagement_id}.docx")
        return os.path.join(
            tempfile.gettempdir(),
            f"OPD_Transformation_Roadmap_{self.engagement_id}.docx"
        )

    # ------------------------------------------------------------------
    # Document assembly
    # ------------------------------------------------------------------

    def _build(self, doc, eng, findings, roadmap, signals, patterns, narrative: dict):
        firm_name = eng.get('firm_name') or self.engagement_id

        self._populate_content_controls(doc, firm_name)
        self._add_cover_page_restriction(doc)

        # Lookup dicts used across multiple sections
        findings_by_id = {f['finding_id']: f for f in findings}
        roadmap_by_id  = {
            item['item_id']: item.get('initiative_name', '')
            for item in roadmap if item.get('item_id')
        }
        initiative_details = {
            d['item_id']: d
            for d in narrative.get('initiative_details', [])
            if isinstance(d, dict) and d.get('item_id')
        }

        # Executive Briefing — unnumbered standalone page, before all numbered sections
        self._build_executive_briefing(doc, findings, narrative, roadmap_by_id)

        # 1 — Executive Summary (four components — Change 1)
        doc.add_heading('Executive Summary', level=1)

        # Component A — 3-4 sentence opening paragraph (narrator)
        opening = narrative.get('executive_summary_opening', '')
        if opening:
            self._add_narrative_paragraphs(doc, _resolve_initiative_codes(opening, roadmap_by_id))
            doc.add_paragraph()

        # Component B — Key Findings at a Glance (sourced from structured data only)
        kf_rows = []
        margin_trend = narrative.get('margin_trend_brief', '')
        if margin_trend:
            kf_rows.append(('Margin Trend', margin_trend))
        # Revenue at Risk — sourced from structured display fields.
        # Preference order: direct_exposure findings (summed if multiple), then largest
        # of any include_in_executive finding, then placeholder.
        _exec_with_fig = [
            f for f in findings
            if f.get('include_in_executive') and f.get('display_figure')
        ]
        _direct = [f for f in _exec_with_fig if f.get('figure_type') == 'direct_exposure']
        _rar_label = 'Revenue at Risk'
        revenue_at_risk = None
        if len(_direct) == 1:
            revenue_at_risk = _direct[0]['display_figure']
        elif len(_direct) > 1:
            _total = sum(
                v for v in (_parse_display_figure_to_float(f['display_figure']) for f in _direct)
                if v is not None
            )
            if _total > 0:
                _fmt = (f'~${_total / 1_000_000:.1f}M' if _total >= 1_000_000
                        else f'~${int(round(_total / 1000))}K')
                revenue_at_risk = f'{_fmt}+'
        elif _exec_with_fig:
            _best = max(_exec_with_fig,
                        key=lambda f: _parse_display_figure_to_float(f['display_figure']) or 0)
            revenue_at_risk = _best['display_figure']
            _rar_label = 'Most urgent active exposure'
        else:
            revenue_at_risk = '[Set in FindingsPanel]'
        kf_rows.append((_rar_label, revenue_at_risk))
        high_findings   = [f for f in findings if f.get('priority') == 'High']
        primary_finding = high_findings[0] if high_findings else (findings[0] if findings else None)
        if primary_finding and primary_finding.get('root_cause'):
            rc  = primary_finding['root_cause']
            dot = rc.find('. ')
            cause_line = (rc[:dot] if dot > 0 else rc).strip()
            # Truncate to 10 words for the at-a-glance box — full explanation in Section 5
            words = cause_line.split()
            if len(words) > 10:
                cause_line = ' '.join(words[:10]) + '\u2026'
            if cause_line:
                kf_rows.append(('Primary Cause', cause_line))
        pz_rows = narrative.get('priority_zero_table_rows', [])
        if pz_rows and isinstance(pz_rows, list) and pz_rows[0].get('action'):
            kf_rows.append(('Act This Week', pz_rows[0]['action']))
        fs_rows = narrative.get('future_state_table_rows', [])
        if fs_rows and isinstance(fs_rows, list):
            # Prefer gross margin row — named metric with current baseline makes the
            # target concrete. Fall back to first row if no gross margin row exists.
            _gm_row = next(
                (r for r in fs_rows
                 if isinstance(r, dict) and 'gross margin' in (r.get('metric') or '').lower()),
                None
            )
            _target_row = _gm_row or (fs_rows[0] if isinstance(fs_rows[0], dict) else None)
            if _target_row and _target_row.get('target'):
                _label = '18-Month Gross Margin Target' if _gm_row else '18-Month Target'
                _current = _target_row.get('current_state', '')
                _val = (_target_row['target'] + f' (currently {_current})'
                        if _current else _target_row['target'])
                kf_rows.append((_label, _val))
        if kf_rows:
            self._key_findings_box(doc, kf_rows)
            doc.add_paragraph()

        # Component C — three short narrative paragraphs (narrator, no CONFIRMED/INFERRED)
        for key in ('executive_summary_para1', 'executive_summary_para2',
                    'executive_summary_para3'):
            para_text = narrative.get(key, '')
            if para_text:
                self._add_narrative_paragraphs(doc, _resolve_initiative_codes(para_text, roadmap_by_id))

        doc.add_paragraph()

        # How to Read This Document — standalone page between Exec Summary and Engagement Overview.
        # Appears in TOC (heading style) but without a section number (w:numPr stripped).
        # Page break before it ensures it starts on its own page.
        self._how_to_read_page(doc, findings, firm_name)
        doc.add_paragraph()

        # 2 — Engagement Overview (redesigned — Change 3)
        doc.add_heading('Engagement Overview', level=1)

        # Component A — compact 3-line metadata block
        for line in [
            f"Client: {firm_name}",
            f"Engagement Date: {eng.get('start_date') or '—'}",
            f"Reference: {self.engagement_id}",
        ]:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Component B — narrator engagement overview paragraph
        overview_para = narrative.get('engagement_overview_paragraph', '')
        if overview_para:
            doc.add_paragraph()
            self._add_narrative_paragraphs(doc, _resolve_initiative_codes(overview_para, roadmap_by_id))
        doc.add_paragraph()

        # Domain Maturity Scorecard
        self._maturity_scorecard(doc, signals, patterns, findings)
        doc.add_paragraph()

        # 3 — Operational Maturity Overview
        doc.add_heading('Operational Maturity Overview', level=1)
        self._signal_table(doc, signals)
        doc.add_paragraph()

        # 4 — Domain Analysis
        doc.add_heading('Domain Analysis', level=1)
        self._findings_by_domain(doc, findings, narrative, roadmap_by_id)
        doc.add_paragraph()

        # 5 — Root Cause Analysis (prose only — finding bullets removed)
        doc.add_heading('Root Cause Analysis', level=1)
        root_cause = narrative.get('root_cause_narrative', '')
        if root_cause:
            self._add_narrative_paragraphs(doc, _resolve_initiative_codes(root_cause, roadmap_by_id))
        else:
            doc.add_paragraph('No root cause narrative generated.')
        doc.add_paragraph()

        # 6 — Economic Impact Analysis (chart + summary table + prose)
        doc.add_heading('Economic Impact Analysis', level=1)

        chart_path = self._generate_economic_chart(findings)
        if chart_path:
            try:
                doc.add_picture(chart_path, width=Inches(6))
                caption = doc.add_paragraph(
                    'Confirmed exposures only. See table below for full economic '
                    'impact including inferred figures.'
                )
                caption.runs[0].italic = True
                caption.runs[0].font.size = Pt(9)
                caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f'Failed to embed economic chart: {e}')
            finally:
                try:
                    os.remove(chart_path)
                except OSError:
                    pass

        self._economic_summary_table(doc, findings)
        doc.add_paragraph()
        econ_narrative = narrative.get('economic_impact_narrative', '')
        if econ_narrative:
            self._add_narrative_paragraphs(doc, _resolve_initiative_codes(econ_narrative, roadmap_by_id))
        doc.add_paragraph()

        # 7 — Future State
        doc.add_heading(f'Where {firm_name} Can Be in 18 Months', level=1)
        future_rows = narrative.get('future_state_table_rows', [])
        if future_rows and isinstance(future_rows, list):
            self._future_state_table(doc, future_rows)
            doc.add_paragraph()
        future_narrative = narrative.get('future_state_narrative', '')
        if future_narrative:
            self._add_narrative_paragraphs(doc, _resolve_initiative_codes(future_narrative, roadmap_by_id))
        doc.add_paragraph()

        # 8 — Transformation Roadmap
        doc.add_heading('Transformation Roadmap', level=1)

        # Visual 2 — Roadmap Timeline chart
        timeline_path = self._generate_roadmap_timeline(roadmap, initiative_details)
        if timeline_path:
            try:
                doc.add_picture(timeline_path, width=Inches(6.5))
                caption = doc.add_paragraph(
                    'Initiative timeline by phase. Months are approximate — '
                    'adjust at kickoff based on available capacity.'
                )
                caption.runs[0].italic = True
                caption.runs[0].font.size = Pt(9)
                caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.add_paragraph()
            except Exception as e:
                logger.warning(f'Failed to embed roadmap timeline: {e}')
            finally:
                try:
                    os.remove(timeline_path)
                except OSError:
                    pass

        # 8.1 — Priority Zero Actions
        doc.add_heading('Priority Zero Actions — Complete This Week', level=2)
        pz_rows = narrative.get('priority_zero_table_rows', [])
        if pz_rows and isinstance(pz_rows, list):
            self._priority_zero_table(doc, pz_rows)
        else:
            doc.add_paragraph('No Priority Zero items identified.')
        doc.add_paragraph()

        # 8.2 — Roadmap Overview
        doc.add_heading('Roadmap Overview', level=2)
        overview_rows = narrative.get('roadmap_overview_rows', [])
        if overview_rows and isinstance(overview_rows, list):
            self._roadmap_overview_table(doc, overview_rows)
        doc.add_paragraph()

        # 8.3 — Quick Wins (High priority + Low effort items, capped at 5)
        quick_wins = [
            r for r in roadmap
            if r.get('priority') == 'High' and r.get('effort') == 'Low'
        ][:5]
        if quick_wins:
            doc.add_heading('Quick Wins — High Priority, Low Effort', level=2)
            doc.add_paragraph(
                'These initiatives deliver significant impact with manageable '
                'implementation effort. They can be started immediately alongside '
                'stabilization work.'
            )
            self._quick_wins_table(doc, quick_wins)
            doc.add_paragraph()

        # 8.4 / 8.5 / 8.6 — Phase Tables
        if roadmap:
            for phase in ['Stabilize', 'Optimize', 'Scale']:
                items = [r for r in roadmap if r.get('phase') == phase]
                if items:
                    doc.add_heading(phase, level=2)
                    rationale = narrative.get('roadmap_rationale', {}).get(phase, '')
                    if rationale:
                        self._add_narrative_paragraphs(doc, _resolve_initiative_codes(rationale, roadmap_by_id))
                        doc.add_paragraph()
                    self._roadmap_phase_table(doc, items, findings_by_id, initiative_details, roadmap_by_id)
                    doc.add_paragraph()
        else:
            doc.add_paragraph('No roadmap items recorded.')

        # 8.6 — Initiative Dependencies
        doc.add_heading('Initiative Dependencies', level=2)
        dep_rows = narrative.get('dependency_table_rows', [])
        if dep_rows and isinstance(dep_rows, list):
            self._dependency_table(doc, dep_rows, roadmap_by_id)
        else:
            doc.add_paragraph('No dependencies identified.')
        doc.add_paragraph()

        # 8.7 — Key Risks
        doc.add_heading('Key Risks', level=2)
        risk_rows = narrative.get('risk_table_rows', [])
        if risk_rows and isinstance(risk_rows, list):
            self._risk_table(doc, risk_rows, roadmap_by_id)
        else:
            doc.add_paragraph('No risks identified in this engagement diagnostic.')
        doc.add_paragraph()

        # 9 — Immediate Next Steps
        doc.add_heading('What Happens Next', level=1)
        doc.add_paragraph(
            'The following actions should be completed within the next two weeks '
            'regardless of any other prioritization decisions.'
        )
        next_rows = narrative.get('next_steps_rows', [])
        if next_rows and isinstance(next_rows, list):
            self._next_steps_table(doc, next_rows)
        else:
            doc.add_paragraph('No next steps generated.')

        doc.add_paragraph()
        self._execution_path_section(
            doc,
            narrative.get('execution_path_recommendation'),
            narrative.get('execution_path_rationale'),
        )

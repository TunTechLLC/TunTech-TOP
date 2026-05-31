"""QA-4 Revision Agent — applies accepted QA items to the v1 roadmap in place.

Architecture (validated by the QA-4 Step 0 model test):
  - Claude (generate_revision_edits) does the judgment: map each accepted QA
    item to a structured edit {type, anchor, context_before, new_text, ...}.
  - This module does the deterministic work: locate each anchor in the loaded
    v1 .docx and apply the edit IN PLACE, preserving every table and style. The
    v1 file is never modified — edits are applied to a freshly loaded copy and
    saved as _v2.docx, so the v1<->v2 diff is exactly the QA contribution.

Matcher (anchor -> document location), in order:
  exact-unique  -> apply
  context-disambiguated (context_before) -> apply
  fuzzy near-miss / not found / multi-paragraph span -> FLAG (never silently
    skip, never risk a mid-word corruption).

`manual`-type edits (new tables, relocations) are recorded as flagged for the
consultant to apply by hand. Everything applied or flagged is persisted to
QARevisionEdits and summarized in a revision note appended to v2.
"""
import os
import bisect
import difflib
import logging

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from api.db.repositories.engagement import EngagementRepository
from api.db.repositories.qa_revision import QARevisionRepository
from api.services.claude import generate_revision_edits
from api.services.qa_inputs import (
    assemble_accepted_qa_items,
    V1_FILENAME_TEMPLATE,
    V2_FILENAME_TEMPLATE,
)

logger = logging.getLogger(__name__)

FUZZY_PREFIX_MIN = 30
FUZZY_RATIO_MIN = 0.85


def reconcile_unaddressed(accepted_items, edits):
    """Return the accepted items that no edit addressed.

    An item is 'addressed' if at least one edit carries its id in
    source_item_id. This closes the silent-omission gap: the v1<->v2 diff can
    only show what changed, never what an accepted item asked for but the agent
    skipped. Reconciliation makes every accepted item provably applied or flagged.
    """
    addressed = {e.get('source_item_id') for e in edits if e.get('source_item_id')}
    return [it for it in accepted_items if it['id'] not in addressed]


# ---------------------------------------------------------------------------
# Document traversal and indexing
# ---------------------------------------------------------------------------

def iter_paragraphs(doc):
    """Yield every paragraph in document order, including table cell paragraphs.

    Mirrors the interleaving that extract_text_from_file produces, so the text
    we send Claude and the text we match anchors against are identical. Nested
    tables (a table inside a cell) are not descended into — rare in roadmaps;
    such content is simply not editable by QA-4 (its anchors flag)."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para


def build_index(paras):
    """Return (texts, starts, full) for a paragraph list.

    full = "\n".join(texts); starts[i] is the char offset of paragraph i in
    full. Used to locate anchors and map offsets back to paragraphs."""
    texts = [p.text for p in paras]
    starts = []
    acc = 0
    for t in texts:
        starts.append(acc)
        acc += len(t) + 1  # +1 for the "\n" separator
    full = "\n".join(texts)
    return texts, starts, full


def _para_index_at(offset, starts, texts):
    """Return the index of the paragraph containing `offset`, or None if the
    offset falls on a separator between paragraphs."""
    if offset < 0:
        return None
    i = bisect.bisect_right(starts, offset) - 1
    if i < 0:
        return None
    if offset < starts[i] + len(texts[i]):
        return i
    return None  # on a "\n" separator


def _all_occurrences(haystack, needle):
    out = []
    start = 0
    while True:
        idx = haystack.find(needle, start)
        if idx == -1:
            break
        out.append(idx)
        start = idx + 1
    return out


def _longest_prefix_len(anchor, full):
    """Longest prefix of `anchor` that occurs in `full` (binary search)."""
    lo, hi, best = 0, len(anchor), 0
    while lo <= hi:
        mid = (lo + hi) // 2
        if mid and anchor[:mid] in full:
            best = mid
            lo = mid + 1
        else:
            hi = mid - 1
    return best


def locate_anchor(anchor, context_before, full):
    """Locate an anchor in `full`. Returns (method, start, end).

    method is one of: 'exact', 'context', 'fuzzy_near_miss', 'ambiguous', 'none'.
    start/end are character offsets into `full` (None when not located)."""
    occ = _all_occurrences(full, anchor)
    if len(occ) == 1:
        return 'exact', occ[0], occ[0] + len(anchor)
    if len(occ) > 1:
        if context_before:
            combo = context_before + anchor
            combo_occ = _all_occurrences(full, combo)
            if len(combo_occ) == 1:
                s = combo_occ[0] + len(context_before)
                return 'context', s, s + len(anchor)
        return 'ambiguous', None, None
    # Not found verbatim — detect a near-miss so the flag is informative.
    plen = _longest_prefix_len(anchor, full)
    if plen >= FUZZY_PREFIX_MIN:
        pos = full.find(anchor[:plen])
        window = full[pos:pos + len(anchor)]
        if difflib.SequenceMatcher(None, anchor, window).ratio() >= FUZZY_RATIO_MIN:
            return 'fuzzy_near_miss', pos, pos + len(window)
    return 'none', None, None


# ---------------------------------------------------------------------------
# In-place edit application
# ---------------------------------------------------------------------------

def replace_in_paragraph(para, anchor, new_text):
    """Replace the first occurrence of `anchor` within a paragraph, across runs,
    preserving the formatting of the run where the anchor begins.

    Returns True if applied, False if the anchor is not found in the paragraph's
    runs (caller flags it)."""
    runs = para.runs
    joined = "".join(r.text for r in runs)
    idx = joined.find(anchor)
    if idx == -1:
        return False
    end = idx + len(anchor)
    pos = 0
    done = False
    for r in runs:
        r_start, r_end = pos, pos + len(r.text)
        if r_end <= idx or r_start >= end:
            pos = r_end
            continue  # run untouched
        left = r.text[:idx - r_start] if idx > r_start else ""
        right = r.text[end - r_start:] if r_end > end else ""
        if not done:
            r.text = left + new_text + right
            done = True
        else:
            r.text = left + right
        pos = r_end
    return True


def insert_paragraph_after(para, text):
    """Insert a new paragraph immediately after `para`, inheriting its style.
    Returns the new Paragraph."""
    new_p = OxmlElement('w:p')
    para._p.addnext(new_p)
    new_para = Paragraph(new_p, para._parent)
    try:
        new_para.style = para.style
    except Exception:
        pass
    new_para.add_run(text)
    return new_para


# ---------------------------------------------------------------------------
# Service
# ---------------------------------------------------------------------------

class QARevisionService:
    """Runs QA-4: load v1.docx, get an edit list from Claude, apply it in place,
    save v2.docx, persist the edits, and append a revision note."""

    def __init__(self, engagement_id: str):
        self.engagement_id = engagement_id
        self.repo = QARevisionRepository()

    def _paths(self, eng: dict):
        reports_folder = eng.get('reports_folder') or ''
        if not reports_folder:
            raise ValueError(
                f"Engagement {self.engagement_id} has no reports_folder configured"
            )
        v1 = os.path.join(reports_folder, V1_FILENAME_TEMPLATE.format(engagement_id=self.engagement_id))
        v2 = os.path.join(reports_folder, V2_FILENAME_TEMPLATE.format(engagement_id=self.engagement_id))
        if not os.path.exists(v1):
            raise FileNotFoundError(
                f"v1 roadmap not found at {v1} — generate the roadmap first"
            )
        return v1, v2

    async def run(self) -> dict:
        """Execute the revision. Returns a summary dict. Raises ValueError /
        FileNotFoundError on missing prerequisites (router maps to 422)."""
        eng = EngagementRepository().get_by_id(self.engagement_id)
        if eng is None:
            raise ValueError(f"Engagement {self.engagement_id} not found")

        v1_path, v2_path = self._paths(eng)
        accepted_block, accepted_items = assemble_accepted_qa_items(self.engagement_id)  # raises if none

        # Load v1 and index it. Send Claude the SAME joined text we match
        # against, so anchors quote exactly what we search.
        doc = Document(v1_path)
        paras = list(iter_paragraphs(doc))
        texts, starts, full = build_index(paras)

        edits = await generate_revision_edits(full, accepted_block)
        if not edits:
            logger.warning(
                f"QA-4 produced no edits for {self.engagement_id} despite accepted "
                f"items — likely an API/parse failure; v2 not written"
            )
            return {
                "engagement_id": self.engagement_id,
                "saved_to": None,
                "edits_total": 0,
                "applied": 0,
                "flagged": 0,
                "manual": 0,
                "note": "No edits were produced — check logs and retry.",
            }

        applied_records = []   # for persistence + note
        for e in edits:
            record = {
                'edit_type': e['type'],
                'qa_source': e['qa_source'],
                'source_item_id': e.get('source_item_id', ''),
                'anchor': e['anchor'],
                'context_before': e['context_before'],
                'new_text': e['new_text'],
                'reason': e['reason'],
                'outcome': None,
                'match_method': None,
            }
            if e['type'] == 'manual':
                record['outcome'] = 'manual'
                applied_records.append(record)
                continue

            method, s, end = locate_anchor(e['anchor'], e['context_before'], full)
            if method in ('exact', 'context'):
                if e['type'] == 'replace':
                    si = _para_index_at(s, starts, texts)
                    ei = _para_index_at(end - 1, starts, texts)
                    if si is None or ei is None or si != ei:
                        record['outcome'] = 'flagged_unresolved'
                        record['match_method'] = 'multiparagraph'
                    elif replace_in_paragraph(paras[si], e['anchor'], e['new_text']):
                        record['outcome'] = 'applied'
                        record['match_method'] = method
                    else:
                        record['outcome'] = 'flagged_unresolved'
                        record['match_method'] = 'anchor_not_in_runs'
                else:  # insert_after
                    ti = _para_index_at(end - 1, starts, texts)
                    if ti is None:
                        ti = _para_index_at(s, starts, texts)
                    if ti is None:
                        record['outcome'] = 'flagged_unresolved'
                        record['match_method'] = 'no_paragraph'
                    else:
                        insert_paragraph_after(paras[ti], e['new_text'])
                        record['outcome'] = 'applied'
                        record['match_method'] = method
            else:
                record['outcome'] = 'flagged_unresolved'
                record['match_method'] = method  # fuzzy_near_miss / ambiguous / none
            applied_records.append(record)

        # Reconcile: any accepted item no edit addressed is recorded as
        # 'unaddressed' so it surfaces on the worklist instead of vanishing.
        for item in reconcile_unaddressed(accepted_items, edits):
            applied_records.append({
                'edit_type': 'unaddressed',
                'qa_source': item['source'],
                'source_item_id': item['id'],
                'anchor': '(not addressed)',
                'context_before': '',
                'new_text': item['summary'],
                'reason': f"Accepted item {item['id']} was not addressed by the revision agent",
                'outcome': 'unaddressed',
                'match_method': None,
            })

        applied      = [r for r in applied_records if r['outcome'] == 'applied']
        flagged      = [r for r in applied_records if r['outcome'] == 'flagged_unresolved']
        manual       = [r for r in applied_records if r['outcome'] == 'manual']
        unaddressed  = [r for r in applied_records if r['outcome'] == 'unaddressed']

        self._append_revision_note(doc, applied, flagged, manual, unaddressed)
        doc.save(v2_path)
        logger.info(
            f"QA-4 saved v2 for {self.engagement_id}: {v2_path} "
            f"(applied {len(applied)}, flagged {len(flagged)}, manual {len(manual)}, "
            f"unaddressed {len(unaddressed)})"
        )

        # Persist: replace any prior revision record for this engagement.
        self.repo.delete_for_engagement(self.engagement_id)
        self.repo.bulk_create(self.engagement_id, applied_records)

        return {
            "engagement_id": self.engagement_id,
            "saved_to": v2_path,
            "v1_path": v1_path,
            "accepted_items": len(accepted_items),
            "edits_total": len(applied_records),
            "applied": len(applied),
            "flagged": len(flagged),
            "manual": len(manual),
            "unaddressed": len(unaddressed),
        }

    def _append_revision_note(self, doc, applied, flagged, manual, unaddressed):
        """Append a QA Revision Note documenting what changed and what still
        needs hand-application."""
        doc.add_heading('QA Revision Note', level=1)
        doc.add_paragraph(
            f"This document (v2) was produced by the TOP QA Revision Agent from the "
            f"v1 roadmap. {len(applied)} edit(s) were applied automatically. "
            f"{len(flagged) + len(manual) + len(unaddressed)} item(s) require manual "
            f"attention (below)."
        )

        if applied:
            doc.add_heading('Applied automatically', level=2)
            for r in applied:
                doc.add_paragraph(
                    f"[{r['qa_source']}] {r['reason'] or r['anchor'][:80]}",
                    style='List Bullet',
                )

        if manual or flagged:
            doc.add_heading('Requires manual application', level=2)
            for r in manual:
                doc.add_paragraph(
                    f"[{r['qa_source']}] STRUCTURAL: {r['reason'] or ''} — {r['new_text'][:200]}",
                    style='List Bullet',
                )
            for r in flagged:
                why = {
                    'fuzzy_near_miss': 'anchor nearly matched but was not verbatim',
                    'ambiguous': 'anchor text appears in more than one place',
                    'multiparagraph': 'change spans a paragraph break',
                    'anchor_not_in_runs': 'anchor could not be located in the paragraph',
                    'no_paragraph': 'insertion point could not be located',
                    'none': 'anchor text was not found in the document',
                }.get(r['match_method'], r['match_method'] or 'could not be located')
                doc.add_paragraph(
                    f"[{r['qa_source']}] {r['reason'] or ''} ({why}). "
                    f"Intended change: \"{r['anchor'][:80]}\" -> \"{r['new_text'][:120]}\"",
                    style='List Bullet',
                )

        if unaddressed:
            doc.add_heading('Accepted items NOT addressed by the agent', level=2)
            doc.add_paragraph(
                "These accepted QA items produced no edit. Review and apply them "
                "by hand — they are not reflected in this document."
            )
            for r in unaddressed:
                doc.add_paragraph(
                    f"[{r['qa_source']} {r['source_item_id']}] {r['new_text'][:200]}",
                    style='List Bullet',
                )

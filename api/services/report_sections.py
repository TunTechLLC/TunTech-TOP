import re
import json
import logging
import tempfile
from collections import defaultdict
from copy import deepcopy

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

logger = logging.getLogger(__name__)

PRIORITY_ORDER = {'High': 0, 'Medium': 1, 'Low': 2}


def _resolve_initiative_codes(text: str, roadmap_by_id: dict | None) -> str:
    """Replace R-code references (e.g. R060, R065) with plain initiative names.
    Shared by the dependency table and key risks table.
    roadmap_by_id: item_id → initiative_name mapping built in _build()."""
    if not roadmap_by_id or not text:
        return text
    return re.sub(r'\bR\d+\b', lambda m: roadmap_by_id.get(m.group(0), m.group(0)), text)


def _shade_cell(cell, hex_color: str):
    """Apply a solid background fill to a table cell. hex_color e.g. 'D9D9D9'."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_inches: list):
    """Set explicit column widths. widths_inches must match table column count."""
    for col_idx, width in enumerate(widths_inches):
        for cell in table.columns[col_idx].cells:
            cell.width = Inches(width)


def _left_align_table(table):
    """Force left alignment on every paragraph in every cell of a table.
    Overrides any justified alignment inherited from the document template."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _strip_economic_source_detail(text: str) -> str:
    """Render-time transform for the Economic Impact cell. Keeps the
    CONFIRMED/DERIVED/INFERRED label but strips calculation detail from
    inside the parentheses.

    DB value is preserved unchanged — downstream processes
    (_parse_economic_figures, structured field suggestion logic) still
    receive the full text.

    Note: regex uses [^)]+ which matches to the first closing paren. Safe
    for current prompt output format which does not generate nested
    parentheses inside the label block. If prompt format changes, revisit
    this regex.
    """
    return re.sub(  # [^)]+ matches to first closing paren — safe unless nested parens appear
        r'\((CONFIRMED|DERIVED|INFERRED):[^)]+\)',
        r'(\1)',
        text,
        flags=re.IGNORECASE,
    )


def _client_facing_evidence(evidence: str) -> str:
    """Transform stored evidence summary to client-facing language.

    Stored: "Supported by P38, P39 across Consulting Economics; 5 signals (3 confirmed, 2 inferred)"
    Client: "Supported by 5 signals across Consulting Economics (3 directly observed, 2 reported)"
    If all signals are High (inferred == 0):
           "Supported by 5 directly observed signals across Consulting Economics"
    Falls back to the original string if the format is not recognised.
    """
    m = re.match(
        r'^Supported by .+ across (.+); (\d+) signals? \((\d+) confirmed, (\d+) inferred\)$',
        evidence,
    )
    if not m:
        return evidence
    domain          = m.group(1)
    total           = int(m.group(2))
    confirmed_count = int(m.group(3))
    inferred_count  = int(m.group(4))
    sig_word = 'signal' if total == 1 else 'signals'
    if inferred_count == 0:
        return f"Supported by {total} directly observed {sig_word} across {domain}"
    return f"Supported by {total} {sig_word} across {domain} ({confirmed_count} directly observed, {inferred_count} reported)"


# -------------------------------------------------------------------
# Domain audience mapping
# -------------------------------------------------------------------

_DOMAIN_AUDIENCE = {
    'AI Readiness':                  'CEO, Director of Operations',
    'Consulting Economics':          'CEO, Finance Lead',
    'Customer Experience':           'CEO, Director of Delivery',
    'Delivery Operations':           'Director of Delivery',
    'Finance and Commercial':        'Finance Lead',
    'Project Governance / PMO':      'Director of Delivery',
    'Resource Management':           'Director of Delivery',
    'Sales & Pipeline':              'VP Sales, CEO',
    'Sales-to-Delivery Transition':  'VP Sales, Director of Delivery',
    'Human Resources':               'Director of Operations',
}

# -------------------------------------------------------------------
# Section map — single source of truth for section numbers.
# Update only here when sections are added, removed, or reordered.
# Only sections referenced by number in the reader guide need entries.
# -------------------------------------------------------------------

_SECTION_MAP = {
    'domain_analysis':   6,
    'root_cause':        7,
    'economic_impact':   8,
    'future_state':      9,
    'roadmap':           10,
    'priority_zero':     '10.1',
    'roadmap_overview':  '10.2',
    # 10.3 Quick Wins is conditionally rendered (High priority + Low effort items only).
    # When it renders, downstream subsections shift by 1. When it is absent they stay
    # at these values. Most engagements will have qualifying items, so the map reflects
    # the common case (Quick Wins present).
    'quick_wins':        '10.3',
    'stabilize':         '10.4',
    'optimize':          '10.5',
    'scale':             '10.6',
    'dependencies':      '10.7',
    'risks':             '10.8',
    'what_happens_next': 11,
}

# -------------------------------------------------------------------
# Reader guide — priority/detail are format strings resolved against
# _SECTION_MAP at render time. Change section numbers in _SECTION_MAP
# only — the guide strings update automatically.
# -------------------------------------------------------------------

_ROLE_READING_GUIDE = [
    {
        'role':            'CEO / Founder',
        'priority':        'Executive Summary, Section {s[priority_zero]}, Section {s[what_happens_next]}',
        'detail':          'Section {s[future_state]} (Future State)',
        'trigger_domains': None,
        'domain_order':    None,
    },
    {
        'role':            'Director of Delivery',
        'priority':        'Section {s[domain_analysis]}{domain_clause}, Section {s[stabilize]}',
        'detail':          'Sections {s[root_cause]}, {s[optimize]}, {s[risks]}',
        'trigger_domains': {
            'Delivery Operations', 'Project Governance / PMO',
            'Resource Management', 'Customer Experience',
        },
        'domain_order':    [
            'Delivery Operations', 'Project Governance / PMO',
            'Resource Management', 'Customer Experience',
        ],
    },
    {
        'role':            'VP Sales / Business Development',
        'priority':        'Section {s[domain_analysis]}{domain_clause}, Section {s[stabilize]}',
        'detail':          'Section {s[dependencies]} (Dependencies)',
        'trigger_domains': {'Sales & Pipeline', 'Sales-to-Delivery Transition'},
        'domain_order':    ['Sales & Pipeline', 'Sales-to-Delivery Transition'],
    },
    {
        'role':            'Finance Lead',
        'priority':        'Section {s[domain_analysis]}{domain_clause}, Section {s[economic_impact]}',
        'detail':          'Sections {s[stabilize]}, {s[optimize]}',
        'trigger_domains': {'Finance and Commercial', 'Consulting Economics'},
        'domain_order':    ['Finance and Commercial', 'Consulting Economics'],
    },
    {
        'role':            'Project Manager / Senior Consultant',
        'priority':        'Section {s[what_happens_next]} (What Happens Next), Section {s[stabilize]}',
        'detail':          'Section {s[risks]} (Key Risks)',
        'trigger_domains': None,
        'domain_order':    None,
    },
    {
        'role':            'Operations / Admin',
        'priority':        'Section {s[domain_analysis]}{domain_clause}, Section {s[optimize]}',
        'detail':          'Section {s[what_happens_next]}',
        'trigger_domains': {'AI Readiness', 'Human Resources'},
        'domain_order':    ['AI Readiness', 'Human Resources'],
    },
]

# -------------------------------------------------------------------
# Interview role / document type derivation
# See CLAUDE.md — "File Naming Convention for OPD Engagements"
# -------------------------------------------------------------------

_INTERVIEW_ROLE_MAP = (
    ('ceo',              'CEO'),
    ('directordelivery', 'Director of Delivery'),
    ('director',         'Director of Delivery'),
    ('vpsales',          'VP of Sales'),
    ('sales',            'VP of Sales'),
    ('financelead',      'Finance Lead'),
    ('finance',          'Finance Lead'),
    ('seniorconsultant', 'Senior Consultant and Project Manager'),
    ('consultant',       'Senior Consultant and Project Manager'),
    ('pm',               'Senior Consultant and Project Manager'),
    ('operations',       'Director of Operations'),
    ('admin',            'Director of Operations'),
)

_DOC_TYPE_MAP = (
    ('financial',      'financial performance documentation'),
    ('portfolio',      'project portfolio summary'),
    ('statusreport',   'project status report'),
    ('status',         'project status report'),
    ('clientfeedback', 'client satisfaction data'),
    ('feedback',       'client satisfaction data'),
    ('sow',            'Statement of Work'),
    ('other',          'supporting documentation'),
)

_FILE_TYPE_FALLBACK = {
    'financial': 'financial performance documentation',
    'sow':       'Statement of Work',
    'status':    'project status report',
    'resource':  'resource planning documentation',
    'delivery':  'delivery documentation',
}


def parse_file_role_and_type(filename: str, file_type: str) -> dict:
    """Parse a single ProcessedFile record into a structured role or document type.

    Kind detection priority:
      1. Interview_ prefix (case-insensitive) → interview
      2. Doc_ prefix (case-insensitive)       → document
      3. file_type == 'interview'             → interview
      4. anything else                        → document

    For interviews: strips _Followup and _N suffixes before role matching.
      Sets is_followup=True when _Followup suffix detected.
    For documents: falls back to _FILE_TYPE_FALLBACK[file_type] when stem
      yields no _DOC_TYPE_MAP match (backward compat for pre-convention files).
    Unrecognised stems are passed through (underscores → spaces) — never 'team member'.

    Returns one of:
      {'kind': 'interview', 'role': str, 'is_followup': bool}
      {'kind': 'document',  'document_type': str}
    """
    stem     = re.sub(r'\.[^.]+$', '', filename)
    name_low = filename.lower()

    if name_low.startswith('interview_'):
        kind     = 'interview'
        raw_stem = stem[10:]
    elif name_low.startswith('doc_'):
        kind     = 'document'
        raw_stem = stem[4:]
    elif file_type == 'interview':
        kind     = 'interview'
        raw_stem = stem
    else:
        kind     = 'document'
        raw_stem = stem

    if kind == 'interview':
        is_followup = bool(re.search(r'_followup', raw_stem, re.I))
        role_stem   = re.sub(r'_followup.*$', '', raw_stem, flags=re.I)
        role_stem   = re.sub(r'_\d+$', '', role_stem)
        role_key    = role_stem.lower().replace('_', '').replace(' ', '')

        for match_key, label in _INTERVIEW_ROLE_MAP:
            if match_key in role_key:
                return {'kind': 'interview', 'role': label, 'is_followup': is_followup}

        return {'kind': 'interview', 'role': role_stem.replace('_', ' '),
                'is_followup': is_followup}

    doc_key = raw_stem.lower().replace('_', '').replace(' ', '')

    for match_key, label in _DOC_TYPE_MAP:
        if match_key in doc_key:
            return {'kind': 'document', 'document_type': label}

    if file_type in _FILE_TYPE_FALLBACK:
        return {'kind': 'document', 'document_type': _FILE_TYPE_FALLBACK[file_type]}

    return {'kind': 'document', 'document_type': raw_stem.replace('_', ' ')}


def _extract_interview_roles(processed_files: list) -> list:
    """Derive interview role list using parse_file_role_and_type().
    Skips followup files (same role, different session). Deduplicates in encounter order."""
    roles = []
    for pf in processed_files:
        result = parse_file_role_and_type(pf.get('file_name', ''), pf.get('file_type', ''))
        if result['kind'] != 'interview':
            continue
        if result.get('is_followup'):
            continue
        role = result['role']
        if role not in roles:
            roles.append(role)
    return roles


def _extract_document_types(processed_files: list) -> list:
    """Derive document type list using parse_file_role_and_type().
    Deduplicates in encounter order."""
    types = []
    for pf in processed_files:
        result = parse_file_role_and_type(pf.get('file_name', ''), pf.get('file_type', ''))
        if result['kind'] != 'document':
            continue
        doc_type = result['document_type']
        if doc_type not in types:
            types.append(doc_type)
    return types


def _compute_confirmed_floor(findings: list) -> str | None:
    """Sum unique confirmed figures across all findings (deduplicates same dollar amount
    appearing in multiple findings). Returns a formatted '~$X.XM+' string or None."""
    seen = set()
    total_val = 0.0
    parseable = False
    for f in sorted(findings, key=lambda x: PRIORITY_ORDER.get(x.get('priority'), 2)):
        conf, _, _ = _parse_economic_figures(f.get('economic_impact', ''))
        if conf == '—':
            continue
        primary = conf.split(', ')[0]
        if primary in seen:
            continue
        seen.add(primary)
        val = _dollar_to_float(primary)
        if val is not None:
            total_val += val
            parseable = True
    if not parseable:
        return None
    if total_val >= 1_000_000:
        return f'~${total_val / 1_000_000:.1f}M+'
    return f'~${int(round(total_val / 1000))}K+'


# Matches dollar amounts including ranges and K/M suffixes, with optional ~ prefix.
_DOLLAR_RE = re.compile(
    r'~?\$[\d,\.]+[KkMmBb]?(?:[–\-]\$?[\d,\.]+[KkMmBb]?)?',
    re.IGNORECASE
)
_CONFIRMED_RE = re.compile(r'\bCONFIRMED(?:-\w+)?', re.IGNORECASE)
_DERIVED_RE   = re.compile(r'\bDERIVED(?:-\w+)?',   re.IGNORECASE)
_INFERRED_RE  = re.compile(r'\bINFERRED(?:-\w+)?',  re.IGNORECASE)


def _is_label_context(clause: str, pos: int, matched_text: str) -> bool:
    """Return True only when a CONFIRMED/DERIVED/INFERRED label at `pos` is being
    used as an economic label rather than as an adjective modifier."""
    if matched_text.isupper():
        return True
    if pos == 0:
        return True
    pre = clause[max(0, pos - 5):pos].rstrip()
    if not pre:
        return True
    return bool(re.search(r'[\(\:\,\-\u2014\u2013]$', pre))


def _dollar_to_float(s: str) -> float | None:
    """Convert a dollar string to a float. Handles K/M suffixes, ~ prefix, commas, ranges.
    Returns None if parsing fails."""
    s = re.sub(r'[~$,\s]', '', s.strip())
    s = re.split(r'[–\-]', s)[0]
    multiplier = 1
    if s and s[-1].upper() == 'K':
        multiplier = 1_000;      s = s[:-1]
    elif s and s[-1].upper() == 'M':
        multiplier = 1_000_000;  s = s[:-1]
    elif s and s[-1].upper() == 'B':
        multiplier = 1_000_000_000; s = s[:-1]
    try:
        return float(s) * multiplier
    except (ValueError, AttributeError):
        return None


def _parse_display_figure_to_float(display_figure: str | None) -> float | None:
    """Parse a display_figure string to float for sorting and summing."""
    if not display_figure:
        return None
    s = display_figure.strip()
    if s.startswith('\u26a0 '):
        s = s[2:].strip()
    m = re.search(r'\s*\u2014', s)
    if m:
        s = s[:m.start()].strip()
    return _dollar_to_float(s)


def _parse_economic_figures(text: str):
    """Parse an economic_impact string into (confirmed, derived, inferred) figure strings.

    Labels must appear AFTER the dollar amount (within 80 chars).
    Label precedence: CONFIRMED > DERIVED > INFERRED.
    Returns '—' for each column when no matching figures are found.
    """
    if not text:
        return '—', '—', '—'

    confirmed_figures = []
    derived_figures   = []
    inferred_figures  = []

    clauses = re.split(r'\.\s+|\.\s*$|;\s*|\n', text)

    for clause in clauses:
        clause = clause.strip()
        if not clause:
            continue

        conf_positions  = [m.start() for m in _CONFIRMED_RE.finditer(clause)
                           if _is_label_context(clause, m.start(), m.group(0))]
        deriv_positions = [m.start() for m in _DERIVED_RE.finditer(clause)
                           if _is_label_context(clause, m.start(), m.group(0))]
        inf_positions   = [m.start() for m in _INFERRED_RE.finditer(clause)
                           if _is_label_context(clause, m.start(), m.group(0))]

        if not conf_positions and not deriv_positions and not inf_positions:
            continue

        for m in _DOLLAR_RE.finditer(clause):
            amt       = m.group(0)
            amt_end   = m.end()
            amt_start = m.start()

            pre_text = clause[:amt_start]
            if pre_text.count('(') > pre_text.count(')'):
                continue

            conf_after  = [p for p in conf_positions  if p >= amt_end]
            deriv_after = [p for p in deriv_positions if p >= amt_end]
            inf_after   = [p for p in inf_positions   if p >= amt_end]

            d_conf  = min((p - amt_end for p in conf_after),  default=9999)
            d_deriv = min((p - amt_end for p in deriv_after), default=9999)
            d_inf   = min((p - amt_end for p in inf_after),   default=9999)

            closest = min(d_conf, d_deriv, d_inf)
            if closest > 80:
                continue

            if d_conf <= d_deriv and d_conf <= d_inf:
                _win_pos = amt_end + d_conf
            elif d_deriv <= d_inf:
                _win_pos = amt_end + d_deriv
            else:
                _win_pos = amt_end + d_inf

            _paren_open  = clause.rfind('(', 0, _win_pos)
            _paren_close = clause.find(')',  _win_pos)
            if (_paren_open != -1 and _paren_close != -1
                    and amt in clause[_paren_open:_paren_close]):
                continue

            if d_conf <= d_deriv and d_conf <= d_inf:
                if amt not in confirmed_figures:
                    confirmed_figures.append(amt)
            elif d_deriv <= d_inf:
                if amt not in derived_figures:
                    derived_figures.append(amt)
            else:
                if amt not in inferred_figures:
                    inferred_figures.append(amt)

    confirmed = ', '.join(confirmed_figures[:3]) if confirmed_figures else '—'
    derived   = ', '.join(derived_figures[:3])   if derived_figures   else '—'
    inferred  = ', '.join(inferred_figures[:3])  if inferred_figures  else '—'
    return confirmed, derived, inferred


# Domain → figure_type mapping for display figure pre-population
_DOMAIN_TO_FIGURE_TYPE = {
    'Consulting Economics':         'annual_drag',
    'Finance and Commercial':       'annual_drag',
    'Sales & Pipeline':             'concentration_risk',
    'Sales-to-Delivery Transition': 'direct_exposure',
    'Delivery Operations':          'direct_exposure',
    'Project Governance / PMO':     'direct_exposure',
    'Resource Management':          'replacement_cost',
    'Customer Experience':          'direct_exposure',
    'AI Readiness':                 'opportunity',
    'Human Resources':              'replacement_cost',
}


def _format_display_figure(raw: str) -> str:
    """Convert a raw extracted dollar string to compact display format.
    "$526,000" -> "$526K", "$1,200,000" -> "$1.2M", "$526K" -> "$526K" (pass-through).
    """
    val = _dollar_to_float(raw)
    if val is None:
        return raw
    if val >= 1_000_000:
        n = val / 1_000_000
        s = f"{n:.1f}"
        if s.endswith('.0'):
            s = s[:-2]
        return f"${s}M"
    if val >= 1_000:
        return f"${val / 1_000:.0f}K"
    return f"${val:.0f}"


def _float_to_dollar(v: float) -> str:
    """Format a float as a compact dollar string. 368000 → '$368K', 1200000 → '$1.2M'."""
    if v >= 1_000_000:
        n = v / 1_000_000
        s = f"{n:.1f}"
        if s.endswith('.0'):
            s = s[:-2]
        return f"${s}M"
    if v >= 1_000:
        return f"${v / 1_000:.0f}K"
    return f"${v:.0f}"


def _prepopulate_display_figure(
    economic_impact_text: str,
    domain: str,
    confirmed_revenue: float | None,
    finding_title: str | None = None,
) -> tuple[str | None, str | None, str | None]:
    """Extract a suggested display_figure, display_label, and figure_type
    from economic_impact text.

    Returns (display_figure, display_label, figure_type) or (None, None, None)
    if no extractable figure exists.

    Guardrail: if the parsed numeric value exceeds confirmed_revenue (when provided),
    prepend a warning prefix to display_figure so the frontend can show a red warning.
    """
    if not economic_impact_text:
        return None, None, None

    confirmed_str, derived_str, inferred_str = _parse_economic_figures(economic_impact_text)

    if confirmed_str != '\u2014':
        raw = confirmed_str.split(', ')[0]
    elif derived_str != '\u2014':
        raw = derived_str.split(', ')[0]
    elif inferred_str != '\u2014':
        raw = inferred_str.split(', ')[0]
    else:
        return None, None, None

    display_figure = _format_display_figure(raw)

    if confirmed_revenue is not None:
        numeric = _dollar_to_float(raw)
        if numeric is not None and numeric > confirmed_revenue:
            display_figure = f'\u26a0 {display_figure}'

    figure_type   = _DOMAIN_TO_FIGURE_TYPE.get(domain, 'direct_exposure')
    display_label = ' '.join(finding_title.split()[:6]) if finding_title else None

    return display_figure, display_label, figure_type


# ===================================================================
# Section renderer mixin
# Inherited by ReportGeneratorService in report_generator.py.
# No imports from report_generator.py — zero circular dependency.
# ===================================================================

class ReportSectionsMixin:
    """Section renderer methods for the OPD Word report.

    All methods receive explicit data parameters — none access self.engagement_id.
    Shared helper functions are module-level in this file.
    """

    # ------------------------------------------------------------------
    # Template helpers
    # ------------------------------------------------------------------

    def _populate_content_controls(self, doc, firm_name: str):
        """Populate named content controls in the template with engagement data."""
        _EXT_PROPS_RT = (
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/'
            'extended-properties'
        )
        _EXT_PROPS_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'

        for rel in doc.part.package.iter_rels():
            if rel.reltype == _EXT_PROPS_RT:
                part = rel.target_part
                root = etree.fromstring(part.blob)
                company_el = root.find(f'{{{_EXT_PROPS_NS}}}Company')
                if company_el is not None:
                    company_el.text = firm_name
                part._blob = etree.tostring(
                    root, xml_declaration=True, encoding='UTF-8', standalone=True
                )
                break

        for sdt in doc.element.body.findall('.//' + qn('w:sdt')):
            alias = sdt.find('.//' + qn('w:alias'))
            if alias is not None and alias.get(qn('w:val')) == 'Company':
                content = sdt.find(qn('w:sdtContent'))
                if content is not None:
                    for t in content.findall('.//' + qn('w:t')):
                        t.text = firm_name
                break

    # ------------------------------------------------------------------
    # Cover page — distribution restriction
    # ------------------------------------------------------------------

    def _add_cover_page_restriction(self, doc):
        """Insert a distribution restriction note on the cover page."""
        RESTRICTION_TEXT = (
            'Distribution: Restricted \u2014 Contains individual performance assessment data. '
            'Distribute only to CEO and Director of Delivery unless performance references '
            'have been reviewed and approved for broader distribution.'
        )
        for para in doc.paragraphs:
            if 'Confidential' in para.text:
                new_p = OxmlElement('w:p')

                pPr_src = para._element.find(qn('w:pPr'))
                if pPr_src is not None:
                    new_p.append(deepcopy(pPr_src))

                new_r = OxmlElement('w:r')
                if para.runs:
                    rPr_src = para.runs[0]._element.find(qn('w:rPr'))
                    if rPr_src is not None:
                        new_r.append(deepcopy(rPr_src))

                new_t = OxmlElement('w:t')
                new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                new_t.text = RESTRICTION_TEXT
                new_r.append(new_t)
                new_p.append(new_r)

                para._element.addnext(new_p)
                logger.info('Distribution restriction note added to cover page')
                return

        logger.warning(
            'Cover page "Confidential" line not found in doc.paragraphs — '
            'distribution restriction note not added (text may be in a text box or drawing object)'
        )

    # ------------------------------------------------------------------
    # Executive Briefing — standalone CEO teaser page
    # ------------------------------------------------------------------

    def _build_executive_briefing(self, doc, findings: list, narrative: dict,
                                  roadmap_by_id: dict | None = None):
        """One-page standalone briefing shown to the CEO before the full report."""
        findings_by_id = {f['finding_id']: f for f in findings if f.get('finding_id')}
        eb = narrative.get('executive_briefing') or {}

        heading_para = doc.add_heading('Executive Briefing', level=1)
        pPr = heading_para._p.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)

        doc.add_paragraph()

        snapshot = eb.get('executive_snapshot', '')
        if snapshot:
            self._add_narrative_paragraphs(doc, _resolve_initiative_codes(snapshot, roadmap_by_id))
            rule = doc.add_paragraph()
            pPr = rule._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            pBdr.append(bottom)
            pPr.append(pBdr)

        raw_problems = eb.get('problems', [])
        problem_rows = []
        for item in (raw_problems[:3] if isinstance(raw_problems, list) else []):
            if not isinstance(item, dict):
                continue
            fid          = item.get('finding_id', '')
            plain_title  = item.get('plain_title', '').strip()
            impact_brief = item.get('impact_brief', '').strip()
            if not plain_title or not impact_brief:
                continue
            if fid not in findings_by_id:
                logger.warning(
                    f"Executive briefing: problem finding_id {fid!r} not in findings — skipping"
                )
                continue
            problem_rows.append((plain_title, impact_brief))

        if problem_rows:
            self._briefing_block_header(doc, 'Three Critical Problems')
            tbl = doc.add_table(rows=0, cols=1)
            tbl.style = 'Table Grid'
            for plain_title, impact_brief in problem_rows:
                row = tbl.add_row()
                cell = row.cells[0]
                title_para = cell.paragraphs[0]
                title_para.clear()
                title_run = title_para.add_run(plain_title)
                title_run.bold = True
                title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                impact_para = cell.add_paragraph()
                impact_para.add_run(impact_brief)
                impact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_col_widths(tbl, [6.5])
            _left_align_table(tbl)

        exec_findings = [
            f for f in findings
            if f.get('include_in_executive')
            and f.get('display_figure')
            and f.get('display_label')
        ]
        exec_findings.sort(
            key=lambda f: (_parse_display_figure_to_float(f['display_figure']) or -1),
            reverse=True,
        )
        number_rows = [
            (f['display_label'], f['display_figure'])
            for f in exec_findings[:3]
        ]

        self._briefing_block_header(doc, 'Three Numbers That Matter')
        if number_rows:
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = 'Table Grid'
            for label, amount in number_rows:
                row = tbl.add_row()
                row.cells[0].text = label
                amt_cell = row.cells[1]
                amt_run  = amt_cell.paragraphs[0].add_run(amount)
                amt_run.bold = True
            _set_col_widths(tbl, [4.5, 2.0])
            _left_align_table(tbl)
        else:
            ph = doc.add_paragraph(
                '[Complete Executive Display fields in FindingsPanel before delivery]'
            )
            if ph.runs:
                ph.runs[0].italic = True
                ph.runs[0].font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
                ph.runs[0].font.size = Pt(9)

        pz_source = narrative.get('priority_zero_table_rows', [])
        action_rows = [
            r for r in (pz_source[:3] if isinstance(pz_source, list) else [])
            if isinstance(r, dict) and r.get('action')
        ]

        if action_rows:
            self._briefing_block_header(doc, 'What Must Happen This Week')
            tbl = doc.add_table(rows=0, cols=1)
            tbl.style = 'Table Grid'
            for r in action_rows:
                row = tbl.add_row()
                row.cells[0].text = r['action']
            _set_col_widths(tbl, [6.5])
            _left_align_table(tbl)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Narrative helpers
    # ------------------------------------------------------------------

    def _briefing_block_header(self, doc, text: str):
        """Styled section anchor for Executive Briefing blocks."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.all_caps  = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_narrative_paragraphs(self, doc, text: str):
        """Add each double-newline-separated paragraph as a Word paragraph, left-aligned."""
        for para in text.split('\n\n'):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ------------------------------------------------------------------
    # Table helpers — existing
    # ------------------------------------------------------------------

    def _kv_table(self, doc, pairs: list):
        """Two-column key-value table. Skips blank values."""
        visible = [(k, v) for k, v in pairs if v and v.strip()]
        if not visible:
            return
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for key, value in visible:
            row = table.add_row()
            row.cells[0].text = key
            row.cells[0].paragraphs[0].runs[0].bold = True
            _shade_cell(row.cells[0], 'F2F2F2')
            row.cells[1].text = value
        _set_col_widths(table, [1.6, 4.9])
        _left_align_table(table)

    def _key_findings_box(self, doc, rows: list):
        """Component B of Executive Summary — Key Findings at a Glance."""
        visible = [(lbl, val) for lbl, val in rows if val]
        if not visible:
            return

        table = doc.add_table(rows=len(visible), cols=2)
        table.style = 'Table Grid'

        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'),   'none')
            b.set(qn('w:sz'),    '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tblBorders.append(b)
        tblPr.append(tblBorders)

        for i, (label, value) in enumerate(visible):
            row = table.rows[i]
            _shade_cell(row.cells[0], 'F2F2F2')
            _shade_cell(row.cells[1], 'F2F2F2')
            row.cells[0].text = label
            if row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].bold = True
            val_str = str(value)
            if val_str.startswith('['):
                run = row.cells[1].paragraphs[0].add_run(val_str)
                run.italic = True
                run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
                run.font.size = Pt(9)
            else:
                row.cells[1].text = val_str

        _set_col_widths(table, [1.6, 4.9])
        _left_align_table(table)

    def _how_to_read_page(self, doc, findings: list, firm_name: str):
        """Prefatory 'How to Read This Document' page."""
        heading_para = doc.add_heading('How to Read This Document', level=1)
        pPr = heading_para._p.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
        pbBefore = OxmlElement('w:pageBreakBefore')
        pbBefore.set(qn('w:val'), 'true')
        pPr.append(pbBefore)

        doc.add_paragraph()

        finding_domains = {f.get('domain', '') for f in findings if f.get('domain')}
        domain_count    = len(finding_domains)
        intro = doc.add_paragraph(
            f"This diagnostic covers findings across {domain_count} operational domains "
            f"and is intended to be used by multiple members of the {firm_name} leadership "
            f"and delivery team. Not every reader needs to read every section. The table "
            f"below identifies which sections are most relevant to each role."
        )
        intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(['If You Are', 'Priority Reading', 'For Full Detail']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [1.5, 3.0, 2.0])

        for entry in _ROLE_READING_GUIDE:
            trigger = entry['trigger_domains']
            if trigger is not None and not (trigger & finding_domains):
                continue

            domain_order = entry.get('domain_order')
            if domain_order:
                active = [d for d in domain_order if d in finding_domains]
                domain_clause = f' ({", ".join(active)})' if active else ''
            else:
                domain_clause = ''

            row = table.add_row()
            row.cells[0].text = entry['role']
            row.cells[1].text = entry['priority'].format(s=_SECTION_MAP, domain_clause=domain_clause)
            row.cells[2].text = entry['detail'].format(s=_SECTION_MAP, domain_clause=domain_clause)
        _left_align_table(table)

    def _signal_table(self, doc, signals: list):
        """Section 3: Operational Maturity Overview."""
        if not signals:
            doc.add_paragraph('No signals recorded.')
            return

        counts = defaultdict(lambda: {'High': 0, 'Medium': 0, 'Hypothesis': 0, 'total': 0})
        for s in signals:
            d = s.get('domain', 'Unknown')
            c = s.get('signal_confidence', 'Medium')
            counts[d][c] = counts[d].get(c, 0) + 1
            counts[d]['total'] += 1

        total_signals = len(signals)
        domain_count  = len(counts)

        intro = (
            f"During this engagement, {total_signals} operational signals were identified "
            f"and reviewed across {domain_count} functional domains. "
            "Signals are classified by the strength of supporting evidence: "
            "Directly Observed signals are confirmed by multiple sources or documented "
            "evidence; Reported signals come from a single source and warrant validation "
            "before acting; Preliminary signals are inferred from indirect evidence and "
            "require further investigation before conclusions can be drawn. "
            "Domains with higher concentrations of Directly Observed signals have the "
            "strongest evidentiary basis for their findings."
        )
        p = doc.add_paragraph(intro)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        for i, h in enumerate(['Domain', 'Signals Reviewed', 'Directly Observed',
                                'Reported', 'Preliminary']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [2.0, 1.0, 1.1, 0.9, 0.9])

        for domain in sorted(counts):
            c = counts[domain]
            row = table.add_row()
            row.cells[0].text = domain
            row.cells[1].text = str(c['total'])
            row.cells[2].text = str(c['High'])
            row.cells[3].text = str(c['Medium'])
            row.cells[4].text = str(c['Hypothesis'])
        _left_align_table(table)

        doc.add_paragraph()
        callout = doc.add_paragraph(
            "Domains with a high concentration of Directly Observed signals have the "
            "strongest evidentiary basis for their findings and are ready for immediate "
            "intervention design. Domains where Preliminary signals predominate may benefit "
            "from additional data collection before final recommendations are finalized."
        )
        callout.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _findings_by_domain(self, doc, findings: list, narrative: dict,
                            roadmap_by_id: dict | None = None):
        """Findings grouped by domain."""
        if not findings:
            doc.add_paragraph('No findings recorded.')
            return

        by_domain = defaultdict(list)
        for f in findings:
            by_domain[f.get('domain', 'Unknown')].append(f)

        domain_analysis = narrative.get('domain_analysis', {})

        for domain in sorted(by_domain):
            doc.add_heading(domain, level=2)

            audience = _DOMAIN_AUDIENCE.get(domain)
            if audience:
                callout     = doc.add_paragraph()
                callout_run = callout.add_run(f'Primary audience: {audience}')
                callout_run.italic          = True
                callout_run.font.size       = Pt(9)
                callout_run.font.color.rgb  = RGBColor(0x66, 0x66, 0x66)

            domain_data = domain_analysis.get(domain, {})
            opening = domain_data.get('opening', '') if isinstance(domain_data, dict) else ''
            closing = domain_data.get('closing', '') if isinstance(domain_data, dict) else ''

            if opening:
                doc.add_paragraph(_resolve_initiative_codes(opening, roadmap_by_id))
                doc.add_paragraph()

            for f in by_domain[domain]:
                doc.add_heading(f.get('finding_title', ''), level=3)
                self._kv_table(doc, [
                    ('Confidence',         f.get('confidence') or ''),
                    ('Priority',           f.get('priority') or ''),
                    ('Effort',             f.get('effort') or ''),
                    ('Operational Impact', f.get('operational_impact') or ''),
                    ('Economic Impact',    _strip_economic_source_detail(f.get('economic_impact') or '')),
                    ('Root Cause',         f.get('root_cause') or ''),
                    ('Recommendation',     f.get('recommendation') or ''),
                ])

                evidence = _client_facing_evidence(f.get('evidence_summary') or '')
                if evidence:
                    ev_para     = doc.add_paragraph()
                    ev_run      = ev_para.add_run(evidence)
                    ev_run.italic         = True
                    ev_run.font.size      = Pt(9)
                    ev_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                raw_quotes = f.get('key_quotes') or ''
                if raw_quotes:
                    try:
                        quotes = json.loads(raw_quotes)
                    except (json.JSONDecodeError, TypeError):
                        quotes = []
                    for quote in quotes:
                        if not quote:
                            continue
                        q_para     = doc.add_paragraph(style='Quote')
                        q_run      = q_para.add_run(f'"{quote}"')
                        q_run.italic    = True
                        q_run.font.size = Pt(9)

                doc.add_paragraph()

            if closing:
                doc.add_paragraph(_resolve_initiative_codes(closing, roadmap_by_id))
                doc.add_paragraph()

    # ------------------------------------------------------------------
    # Table helpers — new sections
    # ------------------------------------------------------------------

    def _generate_economic_chart(self, findings: list):
        """Generate a horizontal bar chart of direct and derived economic exposure by finding."""
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
        except ImportError:
            logger.warning('matplotlib not installed — skipping economic chart')
            return None

        _COLOR_DIRECT  = '#1F3864'
        _COLOR_DERIVED = '#4472C4'

        try:
            chart_data: list = []
            for f in sorted(findings,
                            key=lambda x: PRIORITY_ORDER.get(x.get('priority'), 2)):
                if not f.get('include_in_executive') or not f.get('display_figure'):
                    continue
                val = _parse_display_figure_to_float(f['display_figure'])
                if val is None:
                    continue
                ftype     = f.get('figure_type') or ''
                bar_color = _COLOR_DIRECT if ftype == 'direct_exposure' else _COLOR_DERIVED
                bar_type  = 'Direct' if ftype == 'direct_exposure' else 'Annual Drag'
                raw_title = f.get('finding_title') or ''
                label     = raw_title[:40] + ('...' if len(raw_title) > 40 else '')
                chart_data.append((label, val, bar_color, bar_type))

            if not chart_data:
                return None

            def _fmt(v: float) -> str:
                if v >= 1_000_000:
                    return f'${v / 1_000_000:.1f}M'
                return f'${int(round(v / 1000))}K'

            chart_data.sort(key=lambda x: x[1])
            labels = [d[0] for d in chart_data]
            values = [d[1] for d in chart_data]
            colors = [d[2] for d in chart_data]
            types  = [d[3] for d in chart_data]
            max_val = max(values)

            n = len(chart_data)
            fig, ax = plt.subplots(figsize=(8, max(3.0, n * 0.6)))

            bars = ax.barh(labels, values, color=colors, height=0.5)

            for bar, val in zip(bars, values):
                ax.text(
                    bar.get_width() + max_val * 0.01,
                    bar.get_y() + bar.get_height() / 2,
                    _fmt(val),
                    va='center', ha='left', fontsize=9,
                )

            ax.set_xlabel('Exposure')
            ax.set_xlim(0, max_val * 1.20)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_visible(False)
            ax.yaxis.grid(False)
            ax.xaxis.grid(False)
            ax.tick_params(left=False)
            ax.xaxis.set_major_formatter(
                plt.FuncFormatter(lambda v, _: _fmt(v))
            )
            ax.set_title('Direct and Derived Economic Exposure', fontsize=11, pad=8)

            type_set = set(types)
            if len(type_set) > 1:
                from matplotlib.patches import Patch
                legend_elements = []
                if 'Direct' in type_set:
                    legend_elements.append(Patch(facecolor=_COLOR_DIRECT,  label='Direct Exposure'))
                if 'Annual Drag' in type_set:
                    legend_elements.append(Patch(facecolor=_COLOR_DERIVED, label='Annual Drag'))
                ax.legend(handles=legend_elements, loc='lower right', fontsize=8)

            plt.tight_layout()

            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp_path = tmp.name
            tmp.close()
            fig.savefig(tmp_path, dpi=150, bbox_inches='tight')
            plt.close(fig)
            return tmp_path

        except Exception as exc:
            logger.warning(f'Economic chart generation failed: {exc} — skipping chart')
            return None

    def _generate_roadmap_timeline(self, roadmap: list, initiative_details: dict):
        """Generate a Gantt-style timeline chart for the transformation roadmap."""
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import matplotlib.patches as mpatches
        except ImportError:
            logger.warning('matplotlib not installed — skipping roadmap timeline')
            return None

        if not roadmap:
            return None

        try:
            PHASE_ORDER = {'Stabilize': 0, 'Optimize': 1, 'Scale': 2}
            PHASE_COLORS = {
                'Stabilize': '#2E75B6',
                'Optimize':  '#ED7D31',
                'Scale':     '#70AD47',
            }
            PHASE_BG = {
                'Stabilize': '#EBF3FB',
                'Optimize':  '#FEF3EC',
                'Scale':     '#F0F7ED',
            }
            PHASE_DEFAULTS = {
                'Stabilize': (1, 3),
                'Optimize':  (3, 9),
                'Scale':     (9, 18),
            }

            def _parse_tl(s):
                if not s:
                    return None
                s = s.strip()
                m = re.match(r'Month\s+(\d+)$', s, re.I)
                if m:
                    v = int(m.group(1))
                    return (v, v)
                m = re.match(r'Months?\s+(\d+)\s*[-\u2013]\s*(\d+)$', s, re.I)
                if m:
                    return (int(m.group(1)), int(m.group(2)))
                return None

            rows = []
            for item in roadmap:
                phase    = item.get('phase') or 'Stabilize'
                raw_name = item.get('initiative_name') or ''
                label    = (raw_name[:35] + '\u2026') if len(raw_name) > 35 else raw_name

                detail   = initiative_details.get(item.get('item_id', '')) or {}
                tl       = _parse_tl(detail.get('timeline', ''))
                if tl is None:
                    tl = PHASE_DEFAULTS.get(phase, (1, 18))

                rows.append((label, phase, tl[0], tl[1]))

            if not rows:
                return None

            rows.sort(key=lambda r: (PHASE_ORDER.get(r[1], 9), r[2]))
            rows = list(reversed(rows))
            n = len(rows)

            fig_height = max(3.5, n * 0.46 + 1.2)
            fig, ax    = plt.subplots(figsize=(9, fig_height))

            for phase_name, zone_start, zone_end in [
                ('Stabilize', 0.5, 3),
                ('Optimize',  3,   9),
                ('Scale',     9,   18.5),
            ]:
                ax.axvspan(zone_start, zone_end,
                           color=PHASE_BG.get(phase_name, '#F5F5F5'),
                           alpha=0.6, zorder=0)
                ax.text(
                    (zone_start + zone_end) / 2, 1.01,
                    phase_name,
                    ha='center', va='bottom',
                    fontsize=8, fontweight='bold',
                    color=PHASE_COLORS.get(phase_name, '#555555'),
                    transform=ax.get_xaxis_transform(),
                    clip_on=False,
                )

            for i, (label, phase, start, end) in enumerate(rows):
                width = max(end - start, 0.4)
                ax.barh(i, width, left=start,
                        height=0.55,
                        color=PHASE_COLORS.get(phase, '#888888'),
                        alpha=0.85,
                        zorder=2)

            ax.set_yticks(range(n))
            ax.set_yticklabels([r[0] for r in rows], fontsize=8)
            ax.set_ylim(-0.6, n - 0.4)

            ax.set_xlim(0.5, 18.5)
            ax.set_xticks([1, 3, 6, 9, 12, 15, 18])
            ax.set_xticklabels(
                ['Mo 1', 'Mo 3', 'Mo 6', 'Mo 9', 'Mo 12', 'Mo 15', 'Mo 18'],
                fontsize=8,
            )
            ax.set_xlabel('Timeline (months)', fontsize=9)

            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.tick_params(left=False)
            ax.yaxis.grid(False)
            ax.xaxis.grid(False)

            legend_patches = [
                mpatches.Patch(
                    color=PHASE_COLORS[p], label=p, alpha=0.85
                )
                for p in ['Stabilize', 'Optimize', 'Scale']
                if any(r[1] == p for r in rows)
            ]
            ax.legend(handles=legend_patches, loc='lower right',
                      fontsize=8, frameon=True, framealpha=0.9)

            plt.tight_layout()

            tmp      = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp_path = tmp.name
            tmp.close()
            fig.savefig(tmp_path, dpi=150, bbox_inches='tight')
            plt.close(fig)
            logger.info(f'Roadmap timeline chart generated — {n} initiatives')
            return tmp_path

        except Exception as exc:
            logger.warning(f'Roadmap timeline generation failed: {exc} — skipping chart')
            return None

    def _economic_summary_table(self, doc, findings: list):
        """Section 6 summary table.
        Columns: Finding | Confirmed Exposure | Derived Exposure | Annual Drag (Inferred) | Recovery Potential"""
        rows_with_impact = [
            f for f in sorted(findings, key=lambda x: PRIORITY_ORDER.get(x.get('priority'), 2))
            if f.get('economic_impact')
        ]
        if not rows_with_impact:
            doc.add_paragraph('No economic impact data recorded for this engagement.')
            return

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        for i, h in enumerate(['Finding', 'Confirmed Exposure', 'Derived Exposure',
                                'Annual Drag (Inferred)', 'Recovery Potential']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [1.7, 1.1, 1.1, 1.4, 1.2])

        recovery_map = {'High': 'High', 'Medium': 'Medium', 'Low': 'Low'}

        seen_confirmed   = {}
        seen_derived     = {}
        footnote_markers = {}
        footnotes        = []
        _markers         = ['*', '**', '\u2020', '\u2020\u2020', '\u2021']

        for f in rows_with_impact:
            cf = f.get('confirmed_figure')
            df = f.get('derived_figure')
            af = f.get('annual_drag_figure')
            primary_confirmed = _float_to_dollar(cf) if cf is not None else '—'
            primary_derived   = _float_to_dollar(df) if df is not None else '—'
            primary_inferred  = _float_to_dollar(af) if af is not None else '—'

            display_confirmed = primary_confirmed
            if primary_confirmed != '—':
                if primary_confirmed in seen_confirmed:
                    if primary_confirmed not in footnote_markers:
                        marker = _markers[len(footnote_markers)] if len(footnote_markers) < len(_markers) else '*'
                        footnote_markers[primary_confirmed] = marker
                        footnotes.append((marker, primary_confirmed, seen_confirmed[primary_confirmed], 'confirmed'))
                    display_confirmed = f"{primary_confirmed} \u2014 shared, see note"
                else:
                    seen_confirmed[primary_confirmed] = f.get('finding_title', '')

            display_derived = primary_derived
            if primary_derived != '—':
                if primary_derived in seen_derived:
                    if primary_derived not in footnote_markers:
                        marker = _markers[len(footnote_markers)] if len(footnote_markers) < len(_markers) else '*'
                        footnote_markers[primary_derived] = marker
                        footnotes.append((marker, primary_derived, seen_derived[primary_derived], 'derived'))
                    display_derived = f"{primary_derived} \u2014 shared, see note"
                else:
                    seen_derived[primary_derived] = f.get('finding_title', '')

            row = table.add_row()
            row.cells[0].text = f.get('finding_title') or ''
            row.cells[1].text = display_confirmed
            row.cells[2].text = display_derived
            row.cells[3].text = primary_inferred
            row.cells[4].text = recovery_map.get(f.get('priority', ''), '—')

        _qualifying = [
            f for f in findings
            if f.get('include_in_executive')
            and f.get('display_figure')
            and f.get('figure_type') in ('direct_exposure', 'annual_drag')
        ]
        _dedup: dict = {}
        for _f in _qualifying:
            _fig = _f['display_figure']
            if _fig not in _dedup or _f.get('figure_type') == 'direct_exposure':
                _dedup[_fig] = _f
        _deduped = list(_dedup.values())

        if _deduped:
            _total_val = 0.0
            _parseable = False
            for _f in _deduped:
                _val = _parse_display_figure_to_float(_f['display_figure'])
                if _val is not None:
                    _total_val += _val
                    _parseable  = True
            if _parseable:
                _fmt = (f'~${_total_val / 1_000_000:.1f}M' if _total_val >= 1_000_000
                        else f'~${int(round(_total_val / 1000))}K')
                total_display = (f'{_fmt}+ confirmed floor '
                                 f'(direct exposures only — see individual findings)')
            else:
                total_display = '[Set in FindingsPanel]'
        else:
            total_display = '[Set in FindingsPanel]'

        totals_row = table.add_row()
        for cell in totals_row.cells:
            _shade_cell(cell, 'F2F2F2')
        totals_row.cells[0].text = 'Total Identified Exposure'
        if totals_row.cells[0].paragraphs[0].runs:
            totals_row.cells[0].paragraphs[0].runs[0].bold = True
        if total_display.startswith('['):
            _run = totals_row.cells[1].paragraphs[0].add_run(total_display)
            _run.italic = True
            _run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
            _run.font.size = Pt(9)
        else:
            totals_row.cells[1].text = total_display
        totals_row.cells[2].text = '—'
        totals_row.cells[3].text = '—'
        totals_row.cells[4].text = '—'

        _left_align_table(table)

        for marker, figure, first_title, label_type in footnotes:
            fn = doc.add_paragraph(
                f'{marker} {figure} also appears in "{first_title}" — '
                f'these represent the same underlying cost, not separate exposures. Do not sum.'
            )
            fn.runs[0].italic = True

    def _future_state_table(self, doc, rows: list):
        """Section 7 future state metrics table.
        Columns: Metric | Current State | Benchmark | Target"""
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for i, h in enumerate(['Metric', 'Current State', 'Benchmark', 'Target']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [1.5, 1.5, 1.5, 2.0])

        for r in rows:
            if not isinstance(r, dict):
                continue
            row = table.add_row()
            row.cells[0].text = r.get('metric') or ''
            row.cells[1].text = r.get('current_state') or ''
            row.cells[2].text = r.get('benchmark') or ''
            row.cells[3].text = r.get('target') or ''
        _left_align_table(table)

    def _priority_zero_table(self, doc, rows: list):
        """Section 8.1 Priority Zero table.
        Columns: Action | Owner | What This Unblocks"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(['Action', 'Owner', 'What This Unblocks']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [2.5, 1.2, 2.8])

        for r in rows:
            if not isinstance(r, dict):
                continue
            row = table.add_row()
            row.cells[0].text = r.get('action') or ''
            row.cells[1].text = r.get('owner') or ''
            row.cells[2].text = r.get('what_it_unblocks') or ''
        _left_align_table(table)

    def _roadmap_overview_table(self, doc, rows: list):
        """Section 8.2 Roadmap Overview table.
        Columns: Phase | Timeline | Key Outcomes"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(['Phase', 'Timeline', 'Key Outcomes']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [1.0, 1.2, 4.3])

        for r in rows:
            if not isinstance(r, dict):
                continue
            row = table.add_row()
            row.cells[0].text = r.get('phase') or ''
            row.cells[1].text = r.get('timeline') or ''
            outcomes = r.get('key_outcomes', [])
            if isinstance(outcomes, list):
                row.cells[2].text = '\n'.join(f'\u2022 {o}' for o in outcomes if o)
            else:
                row.cells[2].text = str(outcomes)
        _left_align_table(table)

    def _quick_wins_table(self, doc, items: list):
        """Section 10.3 Quick Wins table — High priority, Low effort roadmap items.
        Columns: Initiative | Domain | Estimated Impact"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(['Initiative', 'Domain', 'Estimated Impact']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'E2EFDA')  # light green — visually distinct from other tables
        _set_col_widths(table, [2.5, 1.5, 2.5])

        for item in items:
            row = table.add_row()
            row.cells[0].text = item.get('initiative_name') or ''
            row.cells[1].text = item.get('domain') or ''
            row.cells[2].text = item.get('estimated_impact') or '\u2014'
        _left_align_table(table)

    def _roadmap_phase_table(self, doc, items: list, findings_by_id: dict,
                              initiative_details: dict, roadmap_by_id: dict | None = None):
        """Sections 10.4/10.5/10.6 phase tables.
        Columns: Initiative | Priority | Effort | Owner | Timeline | Success Metric"""
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        headers = ['Initiative', 'Priority', 'Effort', 'Owner', 'Timeline', 'Success Metric']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [2.2, 0.6, 0.6, 0.9, 0.8, 1.4])

        for item in items:
            item_id = item.get('item_id', '')
            details = initiative_details.get(item_id, {})

            row = table.add_row()

            init_cell = row.cells[0]
            init_cell.paragraphs[0].clear()
            name_run = init_cell.paragraphs[0].add_run(item.get('initiative_name') or '')
            name_run.bold = True
            name_run.font.size = Pt(9)

            capability = item.get('capability') or ''
            if capability:
                cap_para = init_cell.add_paragraph()
                cap_run  = cap_para.add_run(f'Capability: {capability}')
                cap_run.italic    = True
                cap_run.font.size = Pt(8)
                cap_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            row.cells[1].text = item.get('priority') or ''
            row.cells[2].text = item.get('effort') or ''
            row.cells[3].text = item.get('owner') or ''
            row.cells[4].text = details.get('timeline', '') or ''
            row.cells[5].text = details.get('success_metric', '') or ''
        _left_align_table(table)

    def _dependency_table(self, doc, rows: list, roadmap_by_id: dict | None = None):
        """Section 8.6 Initiative Dependencies table.
        Columns: Initiative | Depends On"""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        for i, h in enumerate(['Initiative', 'Depends On']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [3.2, 3.3])

        for r in rows:
            if not isinstance(r, dict):
                continue
            row = table.add_row()
            row.cells[0].text = _resolve_initiative_codes(r.get('initiative') or '', roadmap_by_id)
            row.cells[1].text = _resolve_initiative_codes(r.get('depends_on') or '', roadmap_by_id)
        _left_align_table(table)

    def _risk_table(self, doc, rows: list, roadmap_by_id: dict | None = None):
        """Section 8.7 Key Risks table.
        Columns: Risk | Likelihood | Mitigation"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(['Risk', 'Likelihood', 'Mitigation']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [2.5, 0.8, 3.2])

        for r in rows:
            if not isinstance(r, dict):
                continue
            row = table.add_row()
            row.cells[0].text = _resolve_initiative_codes(r.get('risk') or '', roadmap_by_id)
            row.cells[1].text = r.get('likelihood') or ''
            row.cells[2].text = _resolve_initiative_codes(r.get('mitigation') or '', roadmap_by_id)
        _left_align_table(table)

    def _next_steps_table(self, doc, rows: list):
        """Section 9 Immediate Next Steps table.
        Columns: Action | Owner | Completion Criteria"""
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        for i, h in enumerate(['Action', 'Owner', 'Completion Criteria']):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _shade_cell(cell, 'D9D9D9')
        _set_col_widths(table, [2.5, 1.2, 2.8])

        for r in rows[:10]:
            if not isinstance(r, dict):
                continue
            row = table.add_row()
            row.cells[0].text = r.get('action') or ''
            row.cells[1].text = r.get('owner') or ''
            row.cells[2].text = r.get('completion_criteria') or ''
        _left_align_table(table)

    def _execution_path_section(self, doc, recommendation: str | None,
                                 rationale: str | None):
        """Section 11 subsection — How This Gets Implemented.
        Shows all three execution paths. The recommended path has its first sentence
        (narrator rationale) rendered bold; remaining template text follows in normal weight.
        Falls back gracefully when recommendation is None or unrecognised."""
        doc.add_heading('How This Gets Implemented', level=2)
        doc.add_paragraph(
            'The roadmap in this report can be executed through three models. '
            'The appropriate model depends on the firm\'s internal capacity, '
            'leadership bandwidth, and the urgency of active operational risks.'
        )
        doc.add_paragraph()

        rec = (recommendation or '').lower().strip()

        _PATHS = [
            {
                'key': 'internal',
                'title': 'Path 1 — Internal Execution',
                'template': (
                    'The Priority Zero actions require leadership decisions only. '
                    'The Stabilize phase requires process design and governance changes '
                    'that internal leaders can own with clear accountability. '
                    'This path works when leadership bandwidth is confirmed available '
                    'and a named owner can carry each initiative to completion.'
                ),
            },
            {
                'key': 'guided',
                'title': 'Path 2 — Guided Execution',
                'template': (
                    'The client executes. The consultant provides weekly or biweekly '
                    'leadership alignment, roadmap sequencing, and accountability review — '
                    'ensuring the work gets done correctly and in the right order. '
                    'This is the right model for firms without a dedicated transformation '
                    'function, where leadership needs a structured external reference '
                    'point to stay on sequence.'
                ),
            },
            {
                'key': 'partner',
                'title': 'Path 3 — Partner-Supported Execution',
                'template': (
                    'Specific initiatives are staffed through fractional resources — '
                    'fractional PMO, contractor PMs, or finance operations support. '
                    'The consultant architects the solution and directs the resources. '
                    'This path is appropriate when the firm lacks both internal '
                    'transformation capacity and the leadership bandwidth required '
                    'for a guided engagement.'
                ),
            },
        ]

        for path in _PATHS:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(path['title'])
            title_run.bold = True

            content_para = doc.add_paragraph()
            if rec == path['key'] and rationale:
                first = rationale.strip()
                if not first.endswith('.'):
                    first += '.'
                bold_run = content_para.add_run(first + ' ')
                bold_run.bold = True
                content_para.add_run(path['template'])
            else:
                content_para.add_run(path['template'])

            doc.add_paragraph()
